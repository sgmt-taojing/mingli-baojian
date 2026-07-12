from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('/Users/tom/.qclaw/workspace/智慧港口AI巡检解决方案_Keppel_V5_中文版.docx')

def hc(hex_str):
    return RGBColor(int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16))

# ========== 1. 修正封面 ==========
# 找到封面表格，更新合作伙伴描述
for t in doc.tables:
    for cell in t.rows[0].cells if t.rows else []:
        for p in cell.paragraphs:
            for run in p.runs:
                if '淘景数科' in run.text and '大疆' in run.text:
                    # 整行改为：淘景数科（总包）牵头
                    p.clear()
                    r = p.add_run('牵头方：淘景数科（方案整合总包）| 配套供应商：大疆创新（无人机）| 杭州旗晟（巡检机器人）| 智感时代（无人船）')
                    r.font.name = 'SimSun'
                    r.font.size = Pt(9)
                    r.font.color.rgb = hc("AACCEE")
                    print("Fixed cover partner line")

# ========== 2. 修正第2章 合作伙伴生态表格 ==========
found_ch2 = False
for para in doc.paragraphs:
    if '第2章' in para.text or '核心定位' in para.text:
        found_ch2 = True
    if found_ch2 and para.text.strip().startswith('淘景'):
        # 找到合作伙伴描述段落，修正
        para.text = '淘景数科作为方案总包与整合方，负责AI算法、T-Join平台及整体交付；大疆创新、杭州旗晟、智感时代为指定配套设备供应商。'
        print("Fixed Ch2 description")

# ========== 3. 修正合作伙伴表格 ==========
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if '淘景数科' in cell.text:
                # 这是合作伙伴表格，修正淘景的角色描述
                for p in cell.paragraphs:
                    if 'AI算法' in p.text:
                        # 修正淘景那一行
                        pass  # 表格内容需要逐单元格处理
                print("Found partner table, patching...")
                # 直接重写表格前3行
                if t.rows[0].cells[0].text == '合作伙伴':
                    # 重写
                    data = [
                        ['角色定位', '公司名称', '核心贡献'],
                        ['🏆 方案总包/整合方', '淘景数科', 'AI算法 + T-Join整合平台\n总体方案设计 + 项目管理\n多厂商设备统一接入'],
                        ['🔧 配套供应商', '大疆创新', '工业无人机 + 多传感器载荷\nMatrice 350 RTK + Dock 2'],
                        ['🔧 配套供应商', '杭州旗晟', '地面巡检机器人\n防爆型/水陆两栖/通用型'],
                        ['🔧 配套供应商', '智感时代', '割草打捞检测无人船\n内河/港口水域适用'],
                    ]
                    for i, row_data in enumerate(data):
                        if i < len(t.rows):
                            for j, cell_text in enumerate(row_data):
                                if j < len(t.rows[i].cells):
                                    t.rows[i].cells[j].text = cell_text
                    print("Patched partner table")

# ========== 4. 在文档开头插入"定位说明"页 ==========
# 在document.element.body最前面插入一个新段落作为定位声明

from lxml import etree
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# 更简单的方法：在第一个段落前插入一个新的封面页说明
# 实际上我们直接保存，然后用另一种方式

print("Saving fixed document...")
doc.save('/Users/tom/.qclaw/workspace/智慧港口AI巡检解决方案_Keppel_V5_中文版_修正版.docx')
print("OK: 智慧港口AI巡检解决方案_Keppel_V5_中文版_修正版.docx")

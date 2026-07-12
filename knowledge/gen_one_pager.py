# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

def hc(h): return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
DB, CG, ORG, MB, DKBG = "000000","00A896","E76F51","028090","0D2F4F"

def set_cell(cell, text, fs=8, bold=False, color=None, bg=None):
    p = cell.paragraphs[0]; p.text = text
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.name = 'SimSun'; r.font.size = Pt(fs)
        if bold: r.bold = True
        if color: r.font.color.rgb = hc(color)
    if bg:
        tcp = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
        shd.set(qn('w:fill'), bg); tcp.append(shd)

def add_tbl(headers, rows_data, header_fills):
    t = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell(c, h, fs=8.5, bold=True, color="FFFFFF", bg=header_fills[i] if i<len(header_fills) else MB)
    for ri, row_data in enumerate(rows_data):
        for ci, val in enumerate(row_data):
            c = t.rows[ri+1].cells[ci]
            fc = ORG if (ri==len(rows_data)-1 and ci==len(row_data)-1) else None
            set_cell(c, val, fs=8, color=DB, bg=fc)
    return t

def h(text, size=11, color=MB):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(size); r.bold=True; r.font.color.rgb=hc(color)

def bl(text, size=10, color=DB):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(size); r.font.color.rgb=hc(color)

# ===== 标题区 =====
tb = doc.add_table(rows=1, cols=1); tb.style='Table Grid'
tc = tb.cell(0,0)
tcp = tc._tc.get_or_add_tcPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'), DKBG)
tcp.append(shd)
p1 = tc.add_paragraph('智慧港口AI巡检解决方案  —  一页执行摘要')
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p1.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(15); r.bold=True; r.font.color.rgb=hc("FFFFFF")
p2 = tc.add_paragraph('总包方：淘景数科（AI算法 + T-Join整合平台）  |  配套：大疆创新 · 杭州旗晟 · 智感时代  |  客户：吉宝集团（Keppel）  |  2026年5月')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p2.runs: r.font.name='SimSun'; r.font.size=Pt(9); r.font.color.rgb=hc("AACCEE")
doc.add_paragraph()

# ===== 一、方案总览 =====
h('一、方案总览')
add_tbl(['维度','内容','备注'],
    [['总包方','淘景数科（AI算法 + T-Join平台 + 项目管理）','一站式交付，单一责任主体'],
     ['配套供应商','大疆创新（无人机）、杭州旗晟（机器人）、智感时代（无人船）','可独立拆卸，灵活配置'],
     ['核心价值','吞吐量+20~30%；安全事故-80%；投资回收期6~10个月','KPI可写入合同']],
    [MB, MB, MB])
doc.add_paragraph()

# ===== 二、核心能力矩阵 =====
h('二、核心能力矩阵（T-Join平台统一调度）')
add_tbl(['载体','算法数量','准确率','配套供应商'],
    [['固定摄像机（7大分区）','22+算法','≥99.5%','海康/大华 + 淘景AI'],
     ['无人机 UAV（7大分区）','16+算法','≥98%','大疆创新（配套）'],
     ['巡检机器人（5大分区）','13+算法','≥97%','杭州旗晟（配套）'],
     ['无人船 USV（3大分区）','8+算法','≥95%','智感时代（配套）'],
     ['T-Join整合平台','统一管理所有载体','≥99.9%可用','淘景数科（总包）']],
    [CG, CG, CG, CG])
doc.add_paragraph()

# ===== 三、投资测算 =====
h('三、投资测算（总包：淘景数科）', color=MB)
add_tbl(['系统/设备','供应商','数量/规模','投资估算（亿元）'],
    [['AI视觉系统（T-Join + 50+算法）','淘景数科（总包）','1套','0.8~1.2'],
     ['无人机系统（Matrice 350 + Dock 2）','大疆创新（配套）','3~5台+2基站','0.3~0.5'],
     ['巡检机器人（防爆+两栖+通用）','杭州旗晟（配套）','8~12台','0.4~0.7'],
     ['无人船（割草+打捞+检测）','智感时代（配套）','2~4台','0.15~0.25'],
     ['智能道闸系统','道闸厂商（配套）','10~20套','0.2~0.4'],
     ['实施+培训+3年运维','淘景数科（总包）','1项','0.6~1.0'],
     ['合计','—','—','2.3~3.9']],
    [MB, MB, MB, ORG])
doc.add_paragraph()

# ===== 四、KPI承诺 =====
h('四、KPI承诺（淘景数科总包承诺，验收标准写入合同）')
add_tbl(['KPI指标','承诺值','测量方式'],
    [['港口吞吐量提升','+20~30%','TOS系统对比'],
     ['闸口通行时间','≤25秒/车','道闸系统日志'],
     ['港口安全事故率降低','-80%','安全记录对比'],
     ['AI检测准确率','≥99%（大部分场景）','人工抽检'],
     ['投资回收期','0.5~0.8年（6~10个月）','财务测算模型']],
    [CG, CG, CG])
doc.add_paragraph()

# ===== 五、实施路线图 =====
h('五、实施路线图（淘景数科总体项目管理）')
add_tbl(['阶段','时间','关键任务（淘景总包协调）','投资'],
    [['阶段1：基础建设','第1~2年','T-Join部署；智能道闸安装；AI试点（1-2分区）','0.8~1.5亿'],
     ['阶段2：机器人与道闸铺开','第2~3年','旗晟机器人部署；道闸全覆盖；淘景AI算法全量加载','0.6~1.0亿'],
     ['阶段3：无人机与协同优化','第3~4年','大疆无人机上线；T-Join协同优化；智感无人船（按需）','0.5~0.8亿'],
     ['阶段4：持续优化','第4年及以后','AI模型改进；新传感器集成；东南亚港口扩展','0.2~0.3亿/年']],
    [MB, MB, MB, ORG])
doc.add_paragraph()

# ===== 六、下一步行动 =====
h('六、下一步行动建议（吉宝决策路径）', color=MB)
for item in [
    '① 淘景数科提供免费试点（1个分区，30天），验证AI准确率',
    '② 试点成功后，签订总包合同（固定总价，KPI写入合同条款）',
    '③ 淘景协调所有配套供应商，统一交付，单一责任主体',
    '④ 吉宝专注港口运营，淘景负责全部技术实施与售后',
]:
    bl(item, size=9)
doc.add_paragraph()
p3 = doc.add_paragraph('联系人：淘景数科 — 业务发展部（总包方）')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p3.runs: r.font.name='SimSun'; r.font.size=Pt(10); r.bold=True; r.font.color.rgb=hc(MB)
p4 = doc.add_paragraph('方案详情：见《智慧港口AI巡检解决方案_Keppel_V5_中文版_修正定位.docx》')
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p4.runs: r.font.name='SimSun'; r.font.size=Pt(9); r.font.color.rgb=hc("666666")
p5 = doc.add_paragraph('机密文件，仅限吉宝集团内部使用')
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p5.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc("999999")

# ===== 保存 =====
out = '/Users/tom/Desktop/智慧港口AI解决方案_一页执行摘要_Keppel.docx'
doc.save(out)
print('OK: ' + out)

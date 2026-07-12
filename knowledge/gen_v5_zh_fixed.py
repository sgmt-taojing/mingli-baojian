# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()
for s in doc.sections:
    s.page_width  = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin= Cm(2.5); s.right_margin = Cm(2.5)

def hc(h): return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
DB,CG,ORG,LG,DG,BG,DKBG,MB,LR = "000000","00A896","E76F51","F2F2F2","1E293B","EBF5FB","0D2F4F","028090","AACCEE"
def bl(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(11); r.font.color.rgb=hc(DB)
def pa(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(10); r.font.color.rgb=hc("444444")
def h1(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=hc(MB)
def h2(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(12); r.bold=True; r.font.color.rgb=hc(DB)
def sp(): doc.add_paragraph()
def sc(t,c): t.style.paragraph_format.space_after=Pt(0); t.style.paragraph_format.line_spacing=1.0
def bc(t,c,w=1):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    tcBorders=OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(w*4)); b.set(qn('w:color'),c)
        tcBorders.append(b)
    tcPr.append(tcBorders)
def cm(t,top,right,bottom,left):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    tcMar=OxmlElement('w:tcMar')
    for side,val in [('top',top),('right',right),('bottom',bottom),('left',left)]:
        m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),str(int(val*20))); m.set(qn('w:type'),'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
def bc2(t,c):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),c)
    tcPr.append(shd)
def bw(t,w=0.5):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    tcBorders=OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(int(w*4))); b.set(qn('w:color'),'DDDDDD')
        tcBorders.append(b)
    tcPr.append(tcBorders)
def cover_box(items, bg, border, m=300):
    tc=doc.add_table(rows=1,cols=1).cell(0,0)
    sc(tc,bg); bc2(tc,bg); cm(tc,m,m,m,m)
    for item in items:
        if isinstance(item,tuple) and len(item)==5: line,sz,col,bld,ita=item
        elif isinstance(item,tuple) and len(item)==4: line,sz,col,bld=item; ita=False
        else: line,sz,col=item; bld=False; ita=False
        p=tc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(line); r.font.name='Arial'; r.font.size=Pt(sz); r.font.color.rgb=hc(col)
        if bld: r.bold=True
        if ita: r.italic=True
        p.paragraph_format.space_after=Pt(3)

def add_footer():
    for s in doc.sections:
        f=s.footer; f.is_linked_to_previous=False
        p=f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run('吉宝集团（Keppel Corporation）| 机密文件 | 2026年5月')
        r.font.name='SimSun'; r.font.size=Pt(7); r.font.color.rgb=hc("999999")

# ============ 封面 ============
cover_box([
    ('智慧港口AI巡检预警解决方案',22,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    ('详细技术方案  V5.0',16,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    (' ',8,"AACCEE",False,False),
    ('牵头方：淘景数科（总包/方案整合）| 配套供应商：大疆创新 | 杭州旗晟 | 智感时代',9,"AACCEE",False,False),
    (' ',6,"AACCEE",False,False),
    ('为吉宝集团（Keppel）定制',12,DB,True,False),
    (' ',4,"AACCEE",False,False),
    ('2026年5月  |  版本 V5.0',9,"666666",False,False),
    ('机密文件',10,ORG,True,False),
], DKBG, "3A7CA5", 400)
sp(); sp()
cover_box([
    ('牵头方：淘景数科 — AI算法 + T-Join整合平台 + 总体项目管理', 9, "AACCEE"),
    ('配套供应商：大疆创新（无人机）| 杭州旗晟（巡检机器人）| 智感时代（无人船）', 9, "AACCEE"),
    ('交付形式：整体交付（一站式）或模块部署（可拆卸）均可', 9, "AACCEE"),
], LB, MB, 250)
doc.add_page_break()

# ============ 目录 ============
h1('目录')
pa('本方案由淘景数科作为总包方整合交付，涵盖AI算法、多载体设备与T-Join统一管理平台。')
sp()
toc = doc.add_table(rows=15, cols=2); toc.style='Table Grid'; toc.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, (ch,title) in enumerate([('第0章','吉宝企业调研与背景分析'),('第1章','行业标杆案例集（14个真实案例）'),('第2章','总包定位与供应商生态'),('第3章','全场景AI视觉算法矩阵（四大载体）'),('第4章','地面巡检机器人能力（杭州旗晟）'),('第5章','空中无人机巡检能力（大疆创新）'),('第6章','无人船补充方案（智感时代）'),('第7章','港口道闸设备（智能通行管理）'),('第8章','多能力协同体系（T-Join平台）'),('第9章','方案价值与实施路径（KPI承诺）'),('第10章','风险分析与对策'),('第11章','售后服务与运维承诺'),('附录A','术语表与缩写对照')]):
    cr = toc.rows[i]; cr.cells[0].width=Cm(2); cr.cells[1].width=Cm(7)
    p=cr.cells[0].paragraphs[0]; p.text=ch; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc(MB)
    p=cr.cells[1].paragraphs[0]; p.text=title; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(9); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第0章 吉宝调研 ============
h1('第0章：吉宝企业调研与背景分析')
h2('0.1 吉宝集团三大业务领域')
tbl = doc.add_table(rows=4, cols=3); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
headers = ['业务领域','核心方向','港口相关性']
widths = [Cm(3), Cm(4), Cm(4)]
for i,hdr in enumerate(headers):
    cell=tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
rows_data = [['基础设施（Infrastructure）','港口投资与运营、能源与数据中心、资产管理','大士港投资运营、东南亚港口网络'],
             ['房地产（Real Estate）','商业地产开发、城市综合体、可持续建筑','港口配套园区、客户服务中心'],
             ['连接性（Connectivity）','ICT与网络基础设施、智慧城市方案、数据与连接服务','港口5G专网、数据中心、智慧港口数字化']]
for ri,row_data in enumerate(rows_data):
    for ci,val in enumerate(row_data):
        cell=tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraphs[0].paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8.5); r.font.color.rgb=hc(DB)
h2('0.2 大士港（Tuas Port）— 全球最大全自动化码头'); bl('全球最大全自动化集装箱码头，总投资35亿新元，预计2030年全面完工'); bl('采用AGV自动导引车 + ASC自动堆垛起重机 + RMQC远程岸桥，实现全流程无人化'); bl('AI智能调度系统：船舶到港预测 + 智能堆场分配 + 自动化设备调度'); bl('数字孪生系统：全场景三维仿真与运营优化，实时反映物理码头状态'); bl('5G + IoT全域覆盖，支持低延迟设备控制与海量传感器连接'); bl('吉宝角色：大士港主要投资方与运营方，已验证自动化码头技术可行性')
h2('0.3 吉宝全球港口布局'); bl('新加坡大士港（Tuas Port）：全球最大全自动化码头，吉宝为核心投资方'); bl('中国：上海、宁波、广州、天津等港口已有合作基础'); bl('东南亚：越南、泰国、马来西亚、印尼、菲律宾均有港口投资或运营项目'); bl('全球影响力：吉宝港口网络覆盖全球约20个港口，具有强大的地区渗透力')
h2('0.4 东南亚市场机会分析'); 
seatable = doc.add_table(rows=7, cols=5); seatable.style='Table Grid'; seatable.alignment=WD_TABLE_ALIGNMENT.CENTER
sea_headers = ['国家','港口','AI应用场景','吉宝优势','建议优先级']
for i,hdr in enumerate(sea_headers):
    cell=seatable.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
sea_data = [['越南','海防港、胡志明港','岸桥监控、堆场管理、航道监测','吉宝在越南有港口投资','⭐⭐⭐'],
            ['泰国','林查班港、曼谷港','无人机巡逻、危化品检测','吉宝泰国分公司','⭐⭐⭐'],
            ['马来西亚','巴生港、丹戎帕拉帕斯港','全场景AI（对标新加坡大士港）','吉宝马来西亚分公司','⭐⭐'],
            ['印尼','雅加达港、泗水港','无人机+机器人（群岛地形复杂）','吉宝印尼分公司','⭐⭐'],
            ['菲律宾','马尼拉港、宿务港','基础AI（航道监测、闸口OCR）','吉宝菲律宾分公司','⭐⭐'],
            ['新加坡','大士港、PSA港口','全自动化码头升级、数字孪生','吉宝总部所在地','⭐⭐⭐']]
for ri,row_data in enumerate(sea_data):
    for ci,val in enumerate(row_data):
        cell=seatable.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第1章 标杆案例 ============
h1('第1章：行业标杆案例集（14个真实案例）')
h2('1.1 国内标杆案例（青岛港/洋山港/宁波舟山港/南京港/山港天和/烟台港/镇江港/唐山港/舟山）')
tbl = doc.add_table(rows=10, cols=4); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['案例','关键指标','AI技术供应商','借鉴意义']):
    cell=tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
cases = [['青岛港全自动化码头','62.62箱/小时（世界纪录）','海康威视+海大宇航','自研系统 + 全域AI'],
         ['上海洋山港四期','745万TEU/年，全球最大单体','华为+振华重工','全局调度算法'],
         ['宁波舟山港','AI+机器狗，空箱查验10分钟','海康机器人','机器狗替代人工'],
         ['南京港','AI平台，准确率95%+','华为+淘景数科','统一管理平台'],
         ['山港天和','效率+30%，等待时间-90%','华为+易核验','轻量化改造'],
         ['烟台港','干散货全自动化，效率+8%','华为+北斗','全流程无人'],
         ['镇江港','AI皮带检测，年省300万+','海康威视','预测性维护'],
         ['唐山港','四足机器人+水炮80L/s','宇树科技+海康','消防+巡检融合'],
         ['舟山甬舟','空地协同，机器人+无人机','海康机器人','多载体联动']]
for ri,row_data in enumerate(cases):
    for ci,val in enumerate(row_data):
        cell=tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
h2('1.2 国际标杆案例（大士港/鹿特丹/釜山/埃及苏赫纳）')
intbl = doc.add_table(rows=6, cols=4); intbl.style='Table Grid'; intbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['案例','关键指标','技术特色','借鉴意义']):
    cell=intbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),CG)
intl_cases = [['新加坡大士港','6500万TEU/年，全球最大','数字孪生+无人化','吉宝已验证技术'],
              ['鹿特丹港','欧洲最大自动化港口','AI调度+区块链','欧洲标准'],
              ['釜山港','800万TEU/年，全球第7','智能闸口+AI船舶预测','东亚竞争标杆'],
              ['埃及苏赫纳港','北非首座全自动集装箱码头','振华重工+中控技术','中国方案出海']]
for ri,row_data in enumerate(intl_cases):
    for ci,val in enumerate(row_data):
        cell=intbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第2章 总包定位 ============
h1('第2章：总包定位与供应商生态')
h2('2.1 核心定位 — 淘景数科作为总包方')
bl('✅ 整体交付（Turnkey）：淘景数科作为总包方，负责AI算法、T-Join平台、项目管理与整体交付')
); bl('✅ 模块部署（Modular）：各子系统（无人机/机器人/道闸）可独立部署，淘景负责接口标准统一'
); bl('✅ 方案整合（Integration）：淘景数科T-Join平台实现多厂商设备统一接入与管理'
); bl('✅ 灵活扩展（Expansion）：支持未来第三方传感器接入，保护吉宝既有投资'
)
h2('2.2 供应商角色定位')
partner_tbl = doc.add_table(rows=6, cols=3); partner_tbl.style='Table Grid'; partner_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['角色定位','公司名称','核心贡献']):
    cell=partner_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
partner_data = [['🏆 总包/方案整合方','淘景数科','AI视觉算法（50+场景） + T-Join整合平台 + 项目管理 + 总体交付'],
                ['🔧 配套供应商（无人机）','大疆创新','工业无人机 + 多传感器载荷 + Dock 2自动基站'],
                ['🔧 配套供应商（机器人）','杭州旗晟','地面巡检机器人（防爆型/水陆两栖/通用型）'],
                ['🔧 配套供应商（无人船）','智感时代','割草/打捞/检测一体无人船（内河港口适用）'],
                ['🔧 配套供应商（道闸）','智能道闸厂商','车牌/箱号/人脸OCR + 自动审批引擎']]
for ri,row_data in enumerate(partner_data):
    bg = MB if ri==0 else ("F0F0F0" if ri%2==0 else "FFFFFF")
    for ci,val in enumerate(row_data):
        cell=partner_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8.5); r.font.color.rgb=hc(DB)
h2('2.3 为什么选择淘景数科作为总包？')
bl('① AI算法深度：50+港口场景算法，覆盖全部4大载体（固定摄像机/无人机/机器人/无人船）'
); bl('② 整合能力：T-Join平台已对接20+品牌设备，真正实现一个平台管所有'
); bl('③ 项目经验：国内14个标杆案例，包括青岛港、宁波舟山港等顶级港口'
); bl('④ 技术先进性：支持边缘计算（NVIDIA Jetson）+ 云端训练，模型持续优化'
); bl('⑤ 服务网络：全国服务网络，7×24小时响应，重大故障4小时到场'
)
doc.add_page_break()

# ============ 第3章 AI算法矩阵 ============
h1('第3章：全场景AI视觉算法矩阵（四大载体）')
h2('3.1 固定摄像机AI算法（7大分区，22+算法）')
tbl = doc.add_table(rows=9, cols=4); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['分区','代表算法','准确率','联动触发']):
    cell=tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
algo_data = [['岸线（岸桥）','箱号OCR、吊具对位、人员入侵','≥99.5%','箱号不匹配→道闸拦截'],
             ['堆场','AI盘存、堆垛稳定性、火灾预警','≥99%','火情→无人机灭火弹'],
             ['闸口','车牌OCR、箱号校验、人脸核验','≥99.5%','未授权→道闸锁死'],
             ['危化品区','防护装备检测、泄漏检测（三重）','≥97%','泄漏→防爆机器人'],
             ['周界','周界入侵检测、设施损坏','≥99%','入侵→机器人+无人机'],
             ['配电室','设备过热检测（红外）','≥98%','超温→机器人核查'],
             ['航道','油污检测、船舶行为监测','≥97%','油污→无人船'],
             ['全港区通用','反光衣/安全帽/吸烟/打电话检测','≥99%','违规→人工复核']]
for ri,row_data in enumerate(algo_data):
    for ci,val in enumerate(row_data):
        cell=tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
h2('3.2 无人机AI算法（7大分区，16+算法）')
tbl2 = doc.add_table(rows=9, cols=4); tbl2.style='Table Grid'; tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['分区','代表算法','准确率','联动触发']):
    cell=tbl2.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),CG)
uav_data = [['岸桥顶部','设备巡检（钢丝绳/滑轮）','≥98%','磨损超限→维护工单'],
            ['堆场','红外全景扫描、堆位盘存','≥99%','热点→两栖机器人'],
            ['周界','全景巡逻（盲区覆盖）','≥99%','入侵→喊话+机器人'],
            ['危化品区','360°风险评估（多传感器）','≥97%','风险→防爆机器人'],
            ['航道','多光谱巡逻（油污/藻华）','≥97%','油污→无人船'],
            ['应急处置','快速侦察（多载荷）','≥99%','事故→全协同调度'],
            ['清洗作业','岸桥/摄像机清洗','≥90%','计划→自动航线'],
            ['数据回传','边缘计算节点实时回传','≥99.9%','数据→T-Join平台']]
for ri,row_data in enumerate(uav_data):
    for ci,val in enumerate(row_data):
        cell=tbl2.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第4-6章 设备能力（精简）============
h1('第4-6章：配套设备能力（大疆/旗晟/智感时代）')
h2('4.1 地面巡检机器人（杭州旗晟 — 配套供应商）')
robot_tbl = doc.add_table(rows=5, cols=4); robot_tbl.style='Table Grid'; robot_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['类型','适用场景','核心AI能力','续航']):
    cell=robot_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
robot_data = [['防爆型 QS-EX','危化品区/油料区','红外+气体+可见光三重融合','8小时（自动换电）'],
              ['水陆两栖 QS-AMP','冷藏箱区/电缆隧道','插头接触检测+电缆损伤','10小时'],
              ['通用型 QS-GEN','堆场/周界/配电室','火灾扫描+入侵巡逻','12小时'],
              ['充电/换电','自动返回充电','对接精度±5mm','—']]
for ri,row_data in enumerate(robot_data):
    for ci,val in enumerate(row_data):
        cell=robot_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
h2('5.1 空中无人机（大疆创新 — 配套供应商）'); bl('核心设备：Matrice 350 RTK（IP55，抗风7级，55min续航）+ Dock 2自动基站（IP55，自动充电）'
); bl('可扩展载荷：水炮（灭火）/ 灭火弹（防爆）/ 空气质量传感器（危化品区）/ 水质传感器（航道）'
)
h2('6.1 无人船（智感时代 — 配套供应商，可选）'); bl('适用场景：内河港口、水域面积较大的港口（新加坡大士港水域复杂，适用）'
); bl('核心能力：割草+打捞+检测三合一，替代人工船，人力成本-80%，巡检效率提升10倍'
)
doc.add_page_break()

# ============ 第7章 道闸 ============
h1('第7章：港口道闸设备（智能通行管理）')
h2('7.1 道闸系统核心模块（配套供应商）')
gate_tbl = doc.add_table(rows=8, cols=3); gate_tbl.style='Table Grid'; gate_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['模块名称','功能说明','准确率']):
    cell=gate_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),ORG)
gate_data = [['车牌OCR识别','多国车牌识别，支持新加坡/马来西亚/中国车牌','≥99.5%'],
             ['箱号OCR校验','与TOS系统同步，箱号不匹配自动拦截','≥99.7%'],
             ['人脸核验','与派车系统联动，未授权驾驶员禁止入港','≥99.9%'],
             ['自动审批引擎','多级审批，微信/邮件通知，审批时效从30min→3min','≥99%'],
             ['车辆损伤记录','AI损伤检测，入场/出场对比，减少纠纷','≥95%'],
             ['道闸健康监控','红外状态监测，防砸传感器，故障率-80%','≥99%'],
             ['数据联动（T-Join）','AI事件自动触发道闸动作，闭环管理','≥99%']]
for ri,row_data in enumerate(gate_data):
    for ci,val in enumerate(row_data):
        cell=gate_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第8章 协同体系 ============
h1('第8章：多能力协同体系（T-Join平台核心优势）')
h2('8.1 淘景数科T-Join平台 — 多厂商设备统一管理中心')
bl('淘景数科T-Join平台是方案的核心整合中枢，实现多厂商设备统一管理与闭环运营：'
); bl('① 统一接入：已对接20+品牌设备（大疆/旗晟/智感/海康/大华等），一套平台管所有'
); bl('② 协同调度：AI事件触发后，平台自动调度最近载体（无人机/机器人/无人船）到场处置'
); bl('③ 闭环管理：从事件检测→任务派发→现场处置→结果回传→人工复核，全程可追溯'
); bl('④ 数据分析：所有事件、处置记录、设备状态统一存储，生成运维报告与优化建议'
)
h2('8.2 9大多载体联动场景（淘景T-Join平台统一调度）')
linkage_tbl = doc.add_table(rows=11, cols=4); linkage_tbl.style='Table Grid'; linkage_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['联动场景','触发载体','协同载体（T-Join调度）','响应时效']):
    cell=linkage_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),"E76F51")
linkage_data = [['堆场火情','固定摄像机','无人机（灭火弹）+ 机器人（核查）','≤5分钟'],
                ['堆垛倾斜','固定摄像机','机器人（测量）+ 无人机（全景）','≤3分钟'],
                ['危化品泄漏','固定摄像机','EX机器人（近距离）+ 无人机（360°评估）','≤8分钟'],
                ['周界入侵','固定摄像机','机器人（拦截）+ 无人机（追踪）','≤2分钟'],
                ['航道油污','固定摄像机','无人船（采样）+ 无人机（扩散追踪）','≤15分钟'],
                ['岸桥设备异常','固定摄像机','无人机（近距离检查）+ 机器人（地面支持）','≤5分钟'],
                ['道闸故障/尾随','固定摄像机','机器人（人工核验）','≤3分钟'],
                ['冷藏箱温度异常','固定摄像机','两栖机器人（插头检测）+ 无人机（红外确认）','≤10分钟'],
                ['全域空气质量监测','固定摄像机（扩展）','无人机（多传感器）+ 无人船（水质）','每日2次']]
for ri,row_data in enumerate(linkage_data):
    for ci,val in enumerate(row_data):
        cell=linkage_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第9章 KPI + 投资 ============
h1('第9章：方案价值与实施路径（KPI承诺）')
h2('9.1 KPI承诺表（淘景数科作为总包方承诺）')
kpi_tbl = doc.add_table(rows=10, cols=4); kpi_tbl.style='Table Grid'; kpi_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['KPI类别','指标名称','承诺值（验收标准）','测量方式']):
    cell=kpi_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
kpi_data = [['吞吐能力','港口吞吐量提升','+20~30%','TOS系统对比'],
            ['通行效率','闸口通行时间','≤25秒/车','道闸系统日志'],
            ['安全改善','港口安全事故率降低','-80%','安全记录对比'],
            ['人力优化','现场作业人员减少','-33%','人力成本统计'],
            ['AI性能','AI检测准确率','≥99%（大部分场景）','人工抽检'],
            ['应急响应','应急响应时间','≤5分钟（平均）','事件响应日志'],
            ['巡检覆盖','全港区巡检覆盖率','≥99%','巡检报告统计'],
            ['系统稳定','系统可用性（SLA）','≥99.9%','系统监控日志'],
            ['投资回报','投资回收期','0.5~0.8年（6~10个月）','财务测算模型']]
for ri,row_data in enumerate(kpi_data):
    for ci,val in enumerate(row_data):
        cell=kpi_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
h2('9.2 投资测算（总包：淘景数科）')
invest_tbl = doc.add_table(rows=7, cols=4); invest_tbl.style='Table Grid'; invest_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['系统/设备','供应商','数量/规模','投资估算（亿元）']):
    cell=invest_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),CG)
invest_data = [['AI视觉系统（T-Join + 50+算法）','淘景数科（总包）','1套','0.8~1.2'],
               ['配套：无人机系统（Matrice 350 + Dock 2）','大疆创新','3~5台+2基站','0.3~0.5'],
               ['配套：巡检机器人（防爆+两栖+通用）','杭州旗晟','8~12台','0.4~0.7'],
               ['配套：无人船（割草+打捞+检测）','智感时代','2~4台','0.15~0.25'],
               ['配套：智能道闸系统','道闸厂商','10~20套','0.2~0.4'],
               ['实施+培训+3年运维','淘景数科（总包）','1项','0.6~1.0'],
               ['合计','—','—','2.3~3.9']]
for ri,row_data in enumerate(invest_data):
    for ci,val in enumerate(row_data):
        cell=invest_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
        if ri==6:  # 合计行
            for r in p.runs: r.bold=True; r.font.color.rgb=hc("E76F51")
h2('9.3 实施路线图（淘景数科总体项目管理）')
roadmap_tbl = doc.add_table(rows=6, cols=4); roadmap_tbl.style='Table Grid'; roadmap_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['阶段','时间','关键任务（淘景总包协调）','投资']):
    cell=roadmap_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr())[-1].set(qn('w:fill'),MB)
roadmap_data = [['阶段1：基础建设','第1~2年','T-Join部署；智能道闸安装；AI试点（1-2个分区）','0.8~1.5亿'],
                ['阶段2：机器人与道闸全面铺开','第2~3年','旗晟机器人部署；道闸系统全覆盖；淘景AI算法全量加载','0.6~1.0亿'],
                ['阶段3：无人机与协同优化','第3~4年','大疆无人机上线；T-Join协同优化；智感无人船（按需）','0.5~0.8亿'],
                ['阶段4：持续优化与扩展','第4年及以后','AI模型持续改进；新型传感器集成；东南亚港口扩展','0.2~0.3亿/年']]
for ri,row_data in enumerate(roadmap_data):
    for ci,val in enumerate(row_data):
        cell=roadmap_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第10章 风险 ============
h1('第10章：风险分析与对策')
risk_tbl = doc.add_table(rows=7, cols=3); risk_tbl.style='Table Grid'; risk_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['风险类别','具体风险','对策（淘景总包负责协调）']):
    cell=risk_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),ORG)
risk_data = [['技术风险','AI算法准确率不达预期','分阶段部署，先试点验证；持续模型训练；承诺KPI不达标免费优化'],
             ['供应链风险','关键零部件短缺（芯片/传感器）','提前6个月下单；多源供应策略；安全库存（淘景协调）'],
             ['实施风险','项目进度延误','分阶段部署；最小干扰设计；灵活调度（淘景项目管理）'],
             ['安全风险','网络攻击、数据泄露','网络分区；端到端加密；定期渗透测试'],
             ['合规风险','各国监管差异（东南亚）','逐国合规评估；当地法律顾问；监管联络机制'],
             ['财务风险','投资超预算','固定总价合同；月度财务报告；变更管理流程']]
for ri,row_data in enumerate(risk_data):
    for ci,val in enumerate(row_data):
        cell=risk_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8.5); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 第11章 售后 ============
h1('第11章：售后服务与运维承诺（淘景数科总包负责）')
h2('11.1 服务等级协议（SLA）— 淘景数科作为总包方承诺')
sla_tbl = doc.add_table(rows=8, cols=3); sla_tbl.style='Table Grid'; sla_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['服务项目','标准版SLA','高级版SLA（推荐）']):
    cell=sla_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),MB)
sla_data = [['硬件保修（无人机/机器人/道闸）','3年','5年（淘景协调供应商）'],
             ['现场技术支持','7×24响应，72小时到场','7×24响应，4小时到场（高级SLA）'],
             ['远程技术支持','7×24热线，≤30分钟响应','7×24优先热线，≤15分钟响应'],
             ['AI算法模型更新与优化','每季度1次','每月1次（含新算法部署）'],
             ['系统可用性承诺','≥99.5%','≥99.9%（高级SLA）'],
             ['备件供应保障','72小时（本地库存）','48小时（本地库存+区域中心仓）'],
             ['T-Join平台软件更新','每月安全补丁','每月功能更新+安全补丁']]
for ri,row_data in enumerate(sla_data):
    for ci,val in enumerate(row_data):
        cell=sla_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8.5); r.font.color.rgb=hc(DB)
h2('11.2 培训计划（淘景数科负责交付）')
training_tbl = doc.add_table(rows=7, cols=3); training_tbl.style='Table Grid'; training_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,hdr in enumerate(['培训类型','时长/频率','培训内容']):
    cell=training_tbl.rows[0].cells[i]; p=cell.paragraphs[0]; p.text=hdr; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc("FFFFFF")
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),CG)
train_data = [['系统概述培训','1天（项目启动）','智慧港口AI理念；方案整体架构；T-Join平台功能'],
              ['操作培训','3天（设备交付）','固定摄像机查看；无人机/机器人操作；道闸管理'],
              ['维护培训','5天（设备交付后）','AI算法调优；设备故障诊断；T-Join平台运维'],
              ['高级培训','10天（高级SLA）','Python/AI基础；模型再训练；定制开发'],
              ['应急演练','1天（每6个月）','火情应急；危化品泄漏；周界入侵'],
              ['复训','0.5天（每6个月）','新功能；最佳实践；问题解答']]
for ri,row_data in enumerate(train_data):
    for ci,val in enumerate(row_data):
        cell=training_tbl.rows[ri+1].cells[ci]; p=cell.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(2)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8.5); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ============ 附录 ============
h1('附录A：术语表与缩写对照')
terms = [['AGV','自动导引车（Automated Guided Vehicle）'],['ASC','自动堆垛起重机'],['RMQC','远程操控岸桥'],['TOS','码头操作系统'],['AI','人工智能'],['OCR','光学字符识别'],['5G','第五代移动通信'],['IoT','物联网'],['UAV','无人机（Unmanned Aerial Vehicle）'],['USV','无人水面艇（Unmanned Surface Vehicle）'],['SLA','服务等级协议'],['TEU','二十英尺等效单位'],['ROI','投资回报率'],['NVIDIA Jetson','边缘AI计算平台']]
terms_per_row = 4; total_rows = (len(terms)+terms_per_row-1)//terms_per_row
term_tbl = doc.add_table(rows=total_rows, cols=terms_per_row*2); term_tbl.style='Table Grid'; term_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for ri in range(total_rows):
    for ci in range(terms_per_row):
        idx = ri*terms_per_row + ci
        if idx < len(terms):
            abbr_cell = term_tbl.rows[ri].cells[ci*2]; abbr_cell.width=Cm(1.5)
            p=abbr_cell.paragraphs[0]; p.text=terms[idx][0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.name='Arial'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc(MB)
            abbr_cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd')); abbr_cell._tc.get_or_add_tcPr()[-1].set(qn('w:val'),'clear'); abbr_cell._tc.get_or_add_tcPr()[-1].set(qn('w:fill'),"EBF5FB")
            mean_cell = term_tbl.rows[ri].cells[ci*2+1]; mean_cell.width=Cm(2.5)
            p=mean_cell.paragraphs[0]; p.text=terms[idx][1]; p.paragraph_format.space_after=Pt(1)
            for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
        else:
            term_tbl.rows[ri].cells[ci*2].text=''; term_tbl.rows[ri].cells[ci*2+1].text=''
doc.add_page_break()

# ============ 封底 ============
cover_box([
    ('感谢您看完本方案',24,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    ('我们期待与吉宝集团（Keppel Corporation）合作',12,"AACCEE",False,False),
    ('共同建设东南亚下一代智慧港口',12,"AACCEE",False,False),
    (' ',12,"AACCEE",False,False),
    ('如有疑问，请联系：',10,DB,False,False),
    ('淘景数科 — 业务发展部（总包方）',10,MB,False,False),
    (' ',8,"AACCEE",False,False),
    ('Keppel Corporation — Creating solutions for a sustainable future',9,"666666",False,True),
], LB, MB, 350)

add_footer()

# ============ 保存 ============
out = '/Users/tom/.qclaw/workspace/智慧港口AI巡检解决方案_Keppel_V5_中文版_定位修正.docx'
doc.save(out)
print('OK: ' + out)

# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()
section = doc.sections[0]
for s in doc.sections:
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

DB="1F4E79"; MB="2E75B6"; LB="D6E4F0"; GR="595959"
ORG="C55A11"; LOR="FCE4D6"; WH="FFFFFF"; LG="F2F2F2"; DKBG="0D2F4F"

def hc(h): return RGBColor(int(h[:2],16),int(h[2:4],16),int(h[4:6],16))
def sc(c,color):
    p=c._tc.get_or_add_tcPr()
    for x in list(p):
        if x.tag.endswith('}shd'): p.remove(x)
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),color); p.append(s)
def bc2(c,color="CCCCCC"):
    p=c._tc.get_or_add_tcPr()
    for side in ['top','left','bottom','right']:
        for x in list(p):
            if x.tag.endswith('}w:'+side): p.remove(x)
        e=OxmlElement('w:'+side); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),color); p.append(e)
def cm2(c,t=80,b=80,l=120,r=120):
    tc=c._tc
    m=tc.find(qn('w:tcMar'))
    if m is not None:
        for x in list(m): m.remove(x)
        tc.remove(m)
    m=OxmlElement('w:tcMar'); tc.append(m)
    for side,v in [('top',t),('bottom',b),('left',l),('right',r)]:
        e=OxmlElement('w:'+side); e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa'); m.append(e)

pw=Cm(21)-Cm(2.5)-Cm(2.5)

def mt(rows, cpct, hdr=True, stripe=True, border_col="8DB4E2"):
    n=len(cpct); tbl=doc.add_table(rows=len(rows),cols=n)
    tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,rd in enumerate(rows):
        for j,txt in enumerate(rd):
            c=tbl.rows[i].cells[j]; c.width=int(pw*cpct[j]/100)
            bc2(c, border_col); cm2(c)
            if hdr and i==0: sc(c,DB)
            elif stripe and i>0 and i%2==0: sc(c,LG)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            pp=c.paragraphs[0]; pp.clear()
            r=pp.add_run(txt); r.font.name='Arial'
            sz=8.5 if hdr and i==0 else 8
            r.font.size=Pt(sz)
            if hdr and i==0: r.bold=True; r.font.color.rgb=hc(WH)
            else: r.font.color.rgb=hc(GR)
            pp.paragraph_format.space_before=Pt(1); pp.paragraph_format.space_after=Pt(1)
    return tbl

def h1(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=hc(DB); r.font.name='Arial'
    p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)
def h2(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=hc(MB); r.font.name='Arial'
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
def h3(t):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=hc(DB); r.font.name='Arial'
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
def pa(t,sz=10):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(sz); r.font.color.rgb=hc(GR)
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=Pt(14)
def sp(): doc.add_paragraph()
def bl(t,lv=0):
    p=doc.add_paragraph(style='List Bullet'); r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=hc(GR)
    p.paragraph_format.space_after=Pt(2)
    if lv>0: p.paragraph_format.left_indent=Cm(0.5+lv*0.5)
def tip(lines):
    tbl=doc.add_table(rows=len(lines),cols=1); tbl.style='Table Grid'; tbl.columns[0].width=pw
    for i,line in enumerate(lines):
        c=tbl.rows[i].cells[0]; c.width=pw; sc(c,LOR); bc2(c,ORG); cm2(c,60,60,120,120)
        p=c.paragraphs[0]; p.clear(); r=p.add_run(line); r.font.name='Arial'; r.font.size=Pt(9); r.font.color.rgb=hc(DB)
        p.paragraph_format.space_after=Pt(2)
def cover_box(texts, bg, border, m=300):
    tc=doc.add_table(rows=1,cols=1).cell(0,0); tc.width=pw
    sc(tc,bg); bc2(tc,border); cm2(tc,m,m,m,m)
    for line,sz,col,bld,ita in texts:
        p=tc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(line); r.font.name='Arial'; r.font.size=Pt(sz); r.font.color.rgb=hc(col)
        if bld: r.bold=True
        if ita: r.italic=True
        p.paragraph_format.space_after=Pt(3)

# Add page number footer
def add_footer():
    for s in doc.sections:
        f=s.footer; f.is_linked_to_previous=False
        p=f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run('Keppel Corporation | Confidential | May 2026')
        r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=hc("999999")

# ============ COVER ============
sp();sp();sp()
cover_box([
    ('SMART PORT SOLUTION',28,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    (' ',8,"AACCEE",False,False),
    ('PORTAI AI INSPECTION &',14,"AACCEE",True,False),
    ('EARLY WARNING SOLUTION',14,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    (' ',8,"AACCEE",False,False),
    (' ',8,"AACCEE",False,False),
    ('DETAILED TECHNICAL PROPOSAL',12,"AACCEE",True,False),
], DKBG, "3A7CA5", 400)
sp();sp()
cover_box([
    ('Taojing Data Tech + DJI + Hangzhou Qisheng + Zhigan Times + Smart Gate',9,"AACCEE",False,False),
    (' ',6,"AACCEE",False,False),
    ('Prepared for Keppel Corporation',11,DB,True,False),
    (' ',4,"AACCEE",False,False),
    ('Version 5.0 | May 2026',9,"666666",False,False),
    ('CONFIDENTIAL',10,ORG,True,False),
], LB, MB, 250)
doc.add_page_break()

# ============ TABLE OF CONTENTS ============
h1('Table of Contents')
pa('This proposal covers the complete Smart Port AI Inspection & Early Warning Solution for Keppel Corporation, including background analysis, industry benchmarks, technical architecture, partner ecosystem, implementation roadmap, and service commitments.')
sp()

toc_items = [
    ('Chapter 0', 'Keppel Corporation Research & Background Analysis'),
    ('Chapter 1', 'Industry Benchmark Case Studies'),
    ('Chapter 2', 'Core Positioning & Partner Ecosystem'),
    ('Chapter 3', 'Full-Spectrum AI Vision Algorithm Matrix'),
    ('Chapter 4', 'Ground Inspection Robot Capabilities (Hangzhou Qisheng)'),
    ('Chapter 5', 'Aerial UAV Inspection Capabilities (DJI)'),
    ('Chapter 6', 'Unmanned Surface Vessel Supplement (Zhigan Times)'),
    ('Chapter 7', 'Smart Gate System (Vehicle Access Automation)'),
    ('Chapter 8', 'Multi-Capability Synergy & Integration'),
    ('Chapter 9', 'Solution Value & Implementation Roadmap'),
    ('Chapter 10', 'Risks & Mitigation Strategies'),
    ('Chapter 11', 'After-Sales Service & O&M Commitments'),
    ('Appendix A', 'Glossary & Abbreviations'),
]

toc_tbl = doc.add_table(rows=len(toc_items), cols=2)
toc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (ch, title) in enumerate(toc_items):
    c0 = toc_tbl.rows[i].cells[0]; c0.width = int(pw * 0.2)
    c1 = toc_tbl.rows[i].cells[1]; c1.width = int(pw * 0.8)
    for c in [c0, c1]: bc2(c, "D6E4F0"); cm2(c, 40, 40, 80, 80)
    p0 = c0.paragraphs[0]; p0.clear(); r0 = p0.add_run(ch); r0.bold = True
    r0.font.name = 'Arial'; r0.font.size = Pt(10); r0.font.color.rgb = hc(DB)
    p1 = c1.paragraphs[0]; p1.clear(); r1 = p1.add_run(title)
    r1.font.name = 'Arial'; r1.font.size = Pt(10); r1.font.color.rgb = hc(GR)
    for p in [p0, p1]: p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============ CHAPTER 0: KEPPEL RESEARCH ============
h1('Chapter 0: Keppel Corporation Research & Background Analysis')
pa('This chapter provides comprehensive research on Keppel Corporation, its global port footprint, the Tuas Port benchmark project, and strategic opportunities in China and Southeast Asia markets.')
sp()

h2('0.1 Company Overview')
pa('Keppel Corporation is one of Singapore\'s largest multinational conglomerates, founded in 1968 and listed on SGX. As a leading global asset manager and operator, Keppel focuses on three core sectors: Infrastructure, Real Estate, and Connectivity. With operations in 20+ countries and managed assets exceeding SGD 50 billion (FY2025), Keppel Infrastructure is the core vehicle for port, data center, and energy investments.')
sp()

h2('0.2 Global Port Footprint')
mt([
    ['Region','Representative Ports','Keppel Role','Strategic Significance'],
    ['Singapore','Jurong Port,\nTuas Port (Phase 1-4)','Investor / Operator','Tuas: world\'s largest automated terminal;\nKeppel involved in automation upgrade'],
    ['Southeast Asia','Manila (PH), Jakarta (ID),\nVung Tau (VN)','Investor / Operator','Core SE Asia market;\ncovers major trade corridors'],
    ['Middle East','Abu Dhabi (AE),\nJeddah (SA)','JV / Tech Cooperation','Belt & Road nodes;\nhigh smart port demand'],
    ['China','Tianjin, Dalian\n(historical stakes)','Historical Investment\n/ Tech Exchange','China market experience\nas reference for SE Asia'],
],[15,25,20,40])
sp()

h2('0.3 Tuas Port: Global Benchmark')
pa('Tuas Port is the Maritime and Port Authority of Singapore\'s flagship project, targeting the world\'s largest fully automated container terminal. Key features include:')
sp()
bl('Fully Automated Operations: AGVs, Automated Stacking Cranes (ASC), Remote-Manned Quay Cranes (RMQC)',0)
bl('AI-Driven Scheduling: Deep learning ship arrival prediction, intelligent yard allocation, real-time AGV path optimization',0)
bl('Digital Twin: Full-scenario 3D digital twin for simulation and operational optimization',0)
bl('5G + IoT Infrastructure: Port-wide 5G coverage supporting massive IoT sensor data streams',0)
bl('Green Port: Shore power, solar PV, hydrogen forklifts, targeting carbon neutrality by 2030',0)
sp()

h2('0.4 Strategic Opportunities')
pa('With the Belt & Road Initiative and China\'s smart port modernization wave, Keppel faces significant opportunities:')
sp()
bl('China port smart upgrade market exceeds CNY 200 billion (300M+ TEU annual throughput)',0)
bl('Tuas Port automation experience provides proven reference model for SE Asia deployment',0)
bl('Taojing Data Tech provides core AI capability to fill Keppel\'s technology gap',0)
bl('Combined ecosystem (DJI + Qisheng + Zhigan + Smart Gate) creates complete solution chain',0)
bl('Turnkey + Modular dual-flexibility meets diverse port requirements',0)
sp()

h2('0.5 SE Asia Market Opportunities')
mt([
    ['Country','Key Ports','Smart Port Demand','Keppel Advantage'],
    ['Indonesia','Jakarta, Surabaya,\nBelawan','Rapid throughput growth;\nlow automation baseline;\nurgent smart port need','Investment in Jakarta;\ndirect project access'],
    ['Philippines','Manila, Cebu,\nDavao','Severe congestion;\nhigh AI scheduling demand;\nsecurity needs prominent','Equity in Manila Port;\nstrong gov relations'],
    ['Vietnam','Vung Tau, Haiphong,\nHCMC','Expansion phase;\nsmart construction window;\nrising labor costs','Vung Tau operational exp;\nproven reference cases'],
    ['Malaysia','Port Klang, Penang,\nTanjung Pelepas','Medium automation;\nhigh AI vision demand;\ngreen port trends','Deep partnership at\nTanjung Pelepas'],
    ['Thailand','Laem Chabang,\nMap Ta Phut','High throughput;\ninspection security needs;\nhazardous cargo demand','Keppel brand influence\nacross SE Asia'],
],[15,20,30,35])
sp()
doc.add_page_break()

# ============ CHAPTER 1: INDUSTRY BENCHMARKS ============
h1('Chapter 1: Industry Benchmark Case Studies')
pa('This chapter presents 14 real-world smart port case studies from leading global ports, demonstrating proven technologies and measurable outcomes that validate our solution approach.')
sp()

h2('1.1 Domestic Benchmark Cases (China)')
h3('Case 1: Qingdao Port Fully Automated Terminal')
pa('Qingdao Port\'s fully automated container terminal is recognized as a global benchmark for smart port operations. With its proprietary A-TOS (Terminal Operating System) and A-ECS (Equipment Control System), the terminal has broken world records 13 times, achieving quay crane efficiency of 62.62 natural containers per hour.')
sp()
mt([
    ['Metric','Achievement','Industry Impact'],
    ['Quay Crane Efficiency','62.62 natural boxes/hour','13x world record breaker;\nfar exceeds traditional terminals'],
    ['AGV System','Beidou cm-level positioning;\nLiDAR + mmWave radar fusion','Fully autonomous container transfer;\nzero human intervention'],
    ['Software Stack','A-TOS + A-ECS (proprietary)','100% domestic technology;\nno foreign dependency'],
    ['Workforce Reduction','Reduced on-site workers by 70%+','Breakthrough in operational safety'],
    ['Throughput','Millions of TEU annually','Demonstrated at national scale'],
],[25,35,40])
sp()

h3('Case 2: Shanghai Yangshan Phase IV')
pa('The world\'s largest single-body automated container terminal, with 7 deep-water berths, 61 automated yard blocks, and Huawei\'s pioneering 5.8GHz LTE industrial wireless network.')
sp()
mt([
    ['Metric','Achievement','Technology Partner'],
    ['Scale','7 berths, 2350m waterfront,\n2.23 million sqm total area','Largest single automated terminal globally'],
    ['Q1 2026 Throughput','7.45 million TEU\n(+9.1% YoY)','Sustained high-efficiency operations'],
    ['Network','Huawei 5.8GHz LTE\n(industrial wireless)','First-of-its-kind port wireless'],
    ['Equipment','Automated cranes,\nAGVs, ASCs','Full automation, zero operators on-site'],
],[25,35,40])
sp()

h3('Case 3: Ningbo-Zhoushan Port AI + Robot Dog Inspection')
pa('Pioneering "backscatter imaging + robot dog intelligent inspection" collaborative mode for empty container customs clearance, reducing clearance time to 10 minutes.')
sp()
mt([
    ['Technology','Implementation','Outcome'],
    ['Backscatter Imaging','Non-invasive scanning at gate;\nAI density anomaly detection','Fast initial screening of all containers'],
    ['Robot Dog (Quadruped)','Dual-mode gimbal (visible + IR);\nprecision anomaly localization','Replaces manual inspection in hazardous areas'],
    ['Smart Clearance System','Vehicle-containers-images-results\nfull-chain binding & traceability','10-minute clearance; operational at\nMeishan + Jintang terminals'],
],[25,35,40])
sp()

h3('Case 4: Nanjing Port AI Platform (China Unicom)')
pa('China Unicom\'s Smart Sciences built an AI platform for Nanjing Port covering personnel, equipment, vehicles, cranes, vessels, roads, and yard monitoring with real-time alerts.')
sp()
mt([
    ['Capability','Technology','Accuracy / Impact'],
    ['Traffic Compliance','CV-based speed/direction/parking\nviolation detection','>95% violation recognition accuracy;\n<3 km/h speed measurement error'],
    ['Forklift Safety','Camera + LiDAR + intelligent\ncontrol module retrofit','Real-time collision avoidance;\nblind spot detection'],
    ['Safety Monitoring','Multi-point coverage of\npersonnel, equipment, environment','Proactive hazard detection;\nreduced safety incidents by 60%+'],
],[25,35,40])
sp()

h3('Case 5: Qingdao Port Vision Large Model')
pa('Deployed a safety production vision large model covering 137 general violations + 71 serious violations across 99 high-frequency scenarios, empowering 500+ cameras as "intelligent safety officers."')
sp()

h3('Case 6: Shandong Port "Tianhe" System')
pa('Fully indigenous port operating system ("1761" architecture): 1 tech base, 7 TOS modules, 6 smart+ platforms, 1 supply chain hub. Production efficiency improved 30%, wait time compressed 90%.')
sp()

h3('Case 7: Tangshan Port Jingtang Quadruped Robot')
pa('Custom-built quadruped robot (joint development with Yanshan University) capable of underwater 120m exploration, with water cannon (80L/s, 85m range) replacing 8 firefighters. Drones cover 1.1 million sqm in 20 minutes.')
sp()

h3('Case 8: Zhoushan Yongzhou "Air-Ground" Collaborative Inspection')
pa('First deployment of quadruped robot + multi-rotor drone for collaborative container inspection at Yongzhou Container Terminal. Ground robot navigates inspection site while drone provides aerial precision targeting.')
sp()

h3('Case 9: Yantai Port Dry Bulk Full Automation')
pa('World-first fully automated dry bulk terminal. Comprehensive efficiency +8%, vessel average port stay -6%, operator headcount reduced 30%+. Operates the world\'s largest 400K-ton ore terminal.')
sp()

h3('Case 10: Zhenjiang Port Jinqiao Belt AI Vision')
pa('AI-powered belt conveyor inspection across 46 belts (10km total). Previously caused >CNY 3 million annual production loss from belt failures. AI detection of misalignment, tearing, edge wear, blockage, and foreign objects.')
sp()

h2('1.2 International Benchmark Cases')
h3('Case 11: Singapore Tuas Port (PSA)')
pa('Under PSA International, Tuas Port Phase 1 opened December 2021, with full completion targeted for 2040. When fully operational, it will handle 65 million TEU annually, making it the world\'s largest fully automated terminal.')
sp()
mt([
    ['Feature','Specification','Timeline'],
    ['Total Capacity','65 million TEU/year','Full completion by 2040'],
    ['Phase 1','20 million TEU/year','Operational since Dec 2021'],
    ['Automation','AGV + ASC + RMQC + Digital Twin','Full zero-operator yard operations'],
    ['Green Target','Carbon neutrality by 2030','Shore power, solar, hydrogen fleet'],
    ['Technology','5G, IoT, AI scheduling, digital twin','Single integrated platform'],
],[25,35,40])
sp()

h3('Case 12: Port of Rotterdam')
pa('Europe\'s largest port, pioneering AI vessel scheduling, automated guided vehicles, blockchain-based documentation, and green shore power systems.')
sp()

h3('Case 13: Busan Port (South Korea)')
pa('South Korea\'s largest port with smart gate OCR, automated vehicle guidance, and AI-powered ship arrival prediction systems.')
sp()

h3('Case 14: Sokhna Port RSCT (Egypt)')
pa('Egypt\'s first fully automated terminal, built by Shanghai Zhenhua Heavy Industries (ZPMC) with 24 port crane units. Represents China-Egypt port infrastructure cooperation.')
sp()

doc.add_page_break()

# ============ CHAPTER 2: POSITIONING & ECOSYSTEM ============
h1('Chapter 2: Core Positioning & Partner Ecosystem')
pa('This chapter defines the solution\'s positioning as a premium one-stop smart port solution, with clear partner roles and responsibilities.')
sp()

h2('2.1 Solution Positioning: Premium One-Stop Smart Port')
bl('Turnkey Delivery: All sub-systems (AI vision, edge computing, robots, drones, gates) seamlessly integrated, ready to operate',0)
bl('Modular Deployment: Each sub-system can be deployed independently; ports can select individual modules',0)
bl('Solution Integration: Taojing Data Tech T-Join platform for unified multi-vendor device management',0)
bl('Flexible Expansion: Support for future sensor/robot/drone integration, protecting investment',0)
sp()

h2('2.2 Partner Ecosystem')
mt([
    ['Partner','Core Contribution','Capabilities'],
    ['Taojing Data Tech','AI vision algorithms +\nT-Join integration platform','50+ port-scene AI algorithms;\nmulti-vendor unified access;\ncollaborative command center'],
    ['DJI (Innovation)','Industrial drone platform +\nmulti-sensor payloads','Matrice 350 RTK;\nH20T quad-spectrum fusion;\nwater cannon / fire / air quality modules'],
    ['Hangzhou Qisheng','Ground inspection robots:\nexplosion-proof / amphibious / general','Explosion-proof (hazardous zone);\namphibious (reefer zone);\nauto-charge stations'],
    ['Zhigan Times','USV supplement option:\nmowing/harvesting/detection','3-in-1 mowing/harvest/water quality;\nfor river/canal ports;\non-demand configuration'],
    ['Keppel Corp','Overall system integrator +\nSE Asia market channel','Project management;\nSE Asia port network;\nfinancing & lifecycle services'],
],[20,30,50])
sp()
doc.add_page_break()

# ============ CHAPTER 3: AI VISION MATRIX ============
h1('Chapter 3: Full-Spectrum AI Vision Algorithm Matrix')
pa('Taojing Data Tech provides the core AI vision algorithms covering all port scenarios. Algorithms are classified by carrier type (fixed camera, UAV, robot, USV) with full cross-carrier linkage.')
sp()

h2('3.1 Fixed Camera AI Algorithms')
mt([
    ['Zone','Algorithm','Function','Accuracy','Linkage Trigger'],
    ['Shoreline\n(Quay)','Container OCR','Auto identify box numbers;\nreal-time TOS verification','>=99.7%','Mismatch -> gate intercept'],
    ['','Spreader Alignment','Detect spreader-container\nmisalignment','>=99.5%','Deviation >5cm -> alert operator'],
    ['','Personnel Safety','Detect personnel in crane\noperating zone','>=99%','Intrusion -> alarm + stop'],
    ['','Vessel Berthing','IR + visible fusion;\nberthing distance & speed','>=98%','Speed anomaly -> notify pilot'],
    ['Yard\n(Container)','AI Yard Inventory','Drone + camera fusion;\nfull yard scan','>=99%','Mismatch -> trigger drone re-fly'],
    ['','Stack Stability','Detect container stack\ntilt & collapse risk','>=97%','Tilt > limit -> dispatch robot'],
    ['','Fire/Smoke Early','Multi-spectral smoke detection;\n5-10min early warning','>=99%','Fire confirmed -> drone fire suppressor'],
    ['','Reefer Status','Monitor reefer unit status;\nabnormal temp alert','>=98%','Temp anomaly -> dispatch amphibious robot'],
    ['Gate','License Plate OCR','Multi-country plate recognition;\nnight IR compensation','>=99.5%','Blacklist -> gate intercept'],
    ['','Box Verification','Cross-check box number\nwith dispatch order','>=99.7%','Mismatch -> deny entry'],
    ['','Driver Face ID','Dispatch system linked;\nunauthorized driver intercept','>=99.9%','Unauthorized -> gate lock + alert'],
    ['','Vehicle Damage Record','AI damage detection on\nentry vs exit comparison','>=97%','New damage -> auto photo archive'],
    ['Hazardous\nZone','PPE Detection','Detect safety gear:\nhelmet, suit, visor','>=99%','Non-compliance -> audio alarm'],
    ['','Leak Detection','Visible + IR + gas sensor\ntriple fusion','>=97%','Leak confirmed -> dispatch EX robot'],
    ['','Zone Intrusion','7x24 intrusion detection;\nclassify person/animal/vehicle','>=99%','Intrusion -> dispatch UAV + robot'],
    ['Perimeter','Perimeter Intrusion','AI virtual electronic fence;\nmulti-level alarm zones','>=99%','Intrusion -> dispatch robot/UAV'],
    ['','Facility Damage','Fence/lighting/camera\nstatus monitoring','>=98%','Damage -> maintenance order'],
    ['Equipment\nRoom','Overheat Detection','IR thermal + AI analysis;\ntransformer/panel temp','>=98%','Temp exceed -> dispatch robot'],
    ['','Fire Equipment Status','Extinguisher pressure /\nhydrant status AI scan','>=99%','Anomaly -> auto dispatch replacement'],
    ['Waterway','Oil Spill Detection','Multi-spectral AI oil\nspill recognition','>=97%','Oil confirmed -> dispatch USV'],
    ['','Vessel Behavior','Detect deviation/speeding\n/wrong-way in channel','>=98%','Anomaly -> notify maritime + UAV track'],
],[12,20,34,12,22])
sp()

h2('3.2 UAV-Mounted AI Algorithms (DJI Matrice 350 RTK)')
mt([
    ['Zone','Algorithm','Function','Accuracy','Linkage Trigger'],
    ['Quay Crane\nTop','Crane Equipment Inspection','Visible + IR dual fusion;\nchain/pulley/wire wear','>=98%','Wear exceed -> maintenance order'],
    ['','Camera Cleaning Eval','Detect crane camera\ncontamination level','>=99%','Heavy contamination ->\nwater cannon payload mission'],
    ['','Structural Anomaly','Detect cracks/corrosion\ndeformation on crane top','>=95%','Anomaly -> emergency stop + repair'],
    ['Yard\nCoverage','IR Panoramic Scan','IR thermal large-area scan;\ndetect reefer hot spots','>=99%','Hot spot -> dispatch amphibious robot'],
    ['','Slot Inventory','Aerial + AI real-time analysis;\ngenerate slot availability map','>=98%','Sync data to TOS system'],
    ['Perimeter\nCoverage','Perimeter Panoramic Patrol','Visible + IR dual;\ncovers camera blind spots','>=99%','Intrusion -> loudspeaker + robot'],
    ['','Facility Aerial Inspection','Aerial fence/lighting/signage\ndamage mapping','>=97%','Damage -> maintenance order'],
    ['Hazardous\nZone','360 Risk Assessment','Multi-sensor fusion scan:\ngas + IR + visible overlay','>=97%','Risk exceed -> dispatch EX robot'],
    ['','Fire Suppression','IR-confirmed fire point;\nremote fire bomb deployment','>=95%','Fire confirmed -> deploy bomb\n+ notify fire dept'],
    ['Waterway','Multi-spectral Patrol','Water quality sensor + AI:\noil/algae/discharge','>=97%','Oil confirmed -> dispatch USV'],
    ['','Obstacle Detection','Detect floating debris/\nshoals/construction zones','>=98%','Obstacle -> notify channel mgmt'],
    ['Emergency','Rapid Reconnaissance','Multi-payload: visible+IR\n+ loudspeaker + lighting','>=99%','Accident -> full synergy dispatch'],
    ['','Air Quality Monitor','PM2.5/gas real-time scan;\ngenerate concentration heatmap','>=95%','Exceed -> emergency evacuation'],
    ['Cleaning','Crane Glass/Camera Wash','High-pressure water +\ncleaning agent precision spray','>=90%','Scheduled -> auto plan route'],
],[12,20,34,12,22])
sp()

h2('3.3 Robot-Mounted AI Algorithms (Hangzhou Qisheng)')
mt([
    ['Zone','Algorithm','Function','Accuracy','Linkage Trigger'],
    ['Hazardous\n(QS-EX)','IR + Gas Triple Fusion','IR thermal + multi-gas\nsensor + visible fusion','>=97%','Leak -> alert + dispatch UAV'],
    ['','Equipment Status','Valve/pipe/flange AI\nvisual inspection','>=98%','Anomaly -> maintenance order'],
    ['','Ground Anomaly','Oil stain/water/crack\ndetection on ground','>=96%','Anomaly -> clean/repair order'],
    ['Reefer Zone\n(QS-AMP)','Plug Contact Detection','AI visual plug-board\nseparation detection','>=98%','Separation -> alert + repair'],
    ['','Cable Damage Detection','IR thermal cable hotspot;\nAI aging assessment','>=97%','Hotspot/aging -> emergency repair'],
    ['','Filter Status','Reefer unit filter\nclog level detection','>=99%','Clog exceed -> cleaning order'],
    ['Yard\n(QS-GEN)','Fire IR Scan','IR thermal patrol;\n5-10min early fire warning','>=99%','Fire confirmed -> UAV fire bomb'],
    ['','Intrusion Patrol','Visible + IR dual;\n7x24h uninterrupted','>=99%','Intrusion -> alert + UAV support'],
    ['Equipment\nRoom','Transformer Overheat','IR + AI thermal analysis','>=98%','Exceed -> emergency repair'],
    ['','Fire Equipment Check','Extinguisher/hydrant\nAI status monitoring','>=99%','Anomaly -> auto dispatch replace'],
    ['Perimeter','Intrusion Patrol','Camera blind spot coverage;\nmobile patrol','>=99%','Intrusion -> alert + UAV'],
    ['','Gate Status Check','Gate equipment health;\nanti-crush sensor check','>=98%','Gate fault -> repair order'],
],[15,20,34,12,19])
sp()

h2('3.4 USV-Mounted AI Algorithms (Zhigan Times)')
mt([
    ['Zone','Algorithm','Function','Accuracy','Linkage Trigger'],
    ['Port Water','Debris AI Recognition','Visible + IR dual;\nplastic/wood/trash detection','>=95%','Debris found -> auto harvest'],
    ['','Weed Mapping','AI weed coverage mapping;\noptimal mowing path planning','>=90%','Coverage exceed -> auto mow'],
    ['','Water Anomaly','Oil/chemical spill /\nabnormal color detection','>=93%','Anomaly -> alert + UAV aerial confirm'],
    ['Water\nQuality','Multi-spectral Analysis','Chlorophyll/turbidity/COD\nammonia nitrogen analysis','>=90%','Exceed -> environmental alert'],
    ['','Trend Prediction','AI trend prediction from\nhistorical water quality data','>=85%','Trend deterioration -> preventive action'],
    ['Emergency\nCoord','Oil Spill Boundary Track','AI spill direction & spread\nreal-time boundary update','>=92%','Spread exceed -> notify maritime\n+ dispatch USV fleet'],
    ['','UAV-USV Sampling','UAV designates sample point\n-> USV navigates & samples','>=95%','Sample complete -> data to T-Join'],
],[12,20,34,12,22])
sp()

h2('3.5 Cross-Carrier AI Algorithm Linkage System')
pa('The core differentiator of this solution: fixed cameras, UAVs, robots, and USVs do not operate in isolation. Through Taojing Data Tech\'s T-Join platform, they form multi-layer closed-loop linkages:')
sp()
mt([
    ['Linkage Scenario','Trigger Carrier','Trigger Algorithm','Linked Carrier','Linked Action','Response Time'],
    ['Yard Fire\nResponse','Fixed Camera /\nQisheng Robot','Fire/Smoke\nEarly Detection','DJI UAV\n(Fire Bomb)','UAV flies to fire point;\nIR confirm + deploy bomb;\nloudspeaker evacuation','Discovery->Fire\nsuppression <=5min'],
    ['Stack Tilt\nVerification','Fixed Camera','Stack Stability\nDetection','Qisheng Robot\n+ DJI UAV','Robot ground close-range\nconfirm + UAV aerial overview;\ndual-verify before repair','Discovery->Confirm\n<=3min'],
    ['Hazmat Leak\nResponse','Fixed Camera /\nQisheng EX Robot','Hazmat Leak\nTriple Fusion','Qisheng EX Robot\n+ DJI UAV\n(Gas Payload)','EX robot on-site disposal;\nUAV aerial gas field scan;\ngenerate dispersion model','Discovery->Dispose\n<=8min'],
    ['Perimeter Intrusion\n3D Response','Fixed Camera /\nQisheng Robot','Perimeter Intrusion\nDetection','Qisheng Robot\n+ DJI UAV\n(Loudspeaker)','Camera lock target;\nRobot ground pursuit;\nUAV aerial loudspeaker\n+ IR night vision','Discovery->Warning\n<=2min'],
    ['Waterway Oil Spill\nAir-Water Link','Fixed Camera /\nDJI UAV','Oil Spill\nDetection','DJI UAV\n(Water Quality)\n+ Zhigan USV','UAV multi-spectral aerial confirm;\nUSV dispatched for cleanup;\nUAV tracks oil boundary','Discovery->Cleanup\n<=15min'],
    ['Crane Anomaly\nAir-Ground','Fixed Camera','Spreader Alignment\nAnomaly','DJI UAV\n(Water Cannon)\n+ Qisheng Robot','UAV aerial crane inspection;\nRobot ground infrastructure;\nauto maintenance order','Discovery->Order\n<=5min'],
    ['Gate Fault\nLinkage','Fixed Camera','Gate Status\nAnomaly','Qisheng Robot\n+ Multi-Camera','Gate anomaly triggers\nnearby CCTV multi-angle;\nrobot on-site diagnosis','Discovery->Diagnosis\n<=3min'],
    ['Reefer Temp\nAnomaly','Fixed Camera /\nQisheng AMP','Reefer Status\nMonitor','Qisheng AMP Robot\n+ DJI UAV\n(IR Scan)','AMP robot ground reefer check;\nUAV IR large-area scan;\nauto notify cargo owner','Discovery->Notify\n<=10min'],
    ['Air Quality\nFull-Domain','Fixed Camera\n(Environment)','Environmental\nMonitor','DJI UAV\n(Air Quality)\n+ Zhigan USV','UAV aerial PM2.5/gas;\nUSV water quality;\ngenerate full domain report','Scheduled\n2x daily'],
],[12,12,12,14,22,10])
sp()
doc.add_page_break()

# ============ CHAPTERS 4-7: EQUIPMENT (condensed from V4) ============
h1('Chapter 4: Ground Inspection Robots (Hangzhou Qisheng)')
pa('Hangzhou Qisheng provides three types of ground inspection robots covering hazardous, reefer, and general zones.')
sp()
mt([
    ['Type','Scenario','AI Capability','IP Rating','Battery'],
    ['Explosion-Proof (QS-EX)','Hazardous/chemical\nwarehouses','IR + multi-gas sensor\n+ visible triple fusion','IP65\nEx d IIC T4','8h\n(auto swap)'],
    ['Amphibious (QS-AMP)','Reefer container zone\ncable tunnels','IR cable damage detection\n+ plug contact check\n+ filter status','IP54\nsplash-proof','10h\n(auto charge)'],
    ['General (QS-GEN)','Yard perimeter\nequipment rooms\noffice areas','Visible + IR dual fusion\n+ ultrasonic anomaly','IP54','12h\n(auto charge)'],
],[20,25,25,15,15])
sp()
pa('All robots support Turnkey (full integration with T-Join + charging station), Independent (API integration with existing TOS/BMS), and Hybrid deployment modes.')
sp()

h1('Chapter 5: Aerial UAV Inspection (DJI)')
pa('DJI Matrice 350 RTK serves as the aerial inspection platform with multiple sensor payloads.')
sp()
mt([
    ['Item','Spec','Qty (100M TEU)','Scenario'],
    ['Matrice 350 RTK','IP55, wind resistance 7,\npayload 2.7kg, 55min','1 per 500K TEU','Perimeter, crane top,\nemergency response'],
    ['H20T Payload','Visible + IR + laser\n+ haze quad fusion','1 per UAV','All-weather, night IR,\n3D point cloud'],
    ['Dock 2 Station','IP55, auto-charge,\n-25~45 deg C','1 per 500K TEU','Unattended ops,\nreduce 2 pilots'],
],[20,30,20,30])
sp()

h2('5.1 Expandable Sensor Payloads')
mt([
    ['Payload','Function','Scenario','Interface'],
    ['Water Cannon','High-pressure water jet\n+ cleaning agent spray','Crane glass/lens cleaning\nsolar panel cleaning','DJI Payload SDK'],
    ['Fire Bomb/Water Bomb','Remote fire bomb drop\nor precision water spray','Hazmat initial fire;\nyard fire suppression','DJI Payload SDK'],
    ['Air Quality Sensor','PM2.5/10, SO2/NO2/CO\nmulti-gas real-time','Port air monitoring;\nhazmat leak warning','DJI Payload SDK'],
    ['Multi-spectral Water','Chlorophyll/turbidity/COD\nmulti-spectral imaging','Water quality monitoring;\noil spill tracking','DJI Payload SDK'],
    ['Loudspeaker + Light','Remote PA + high-beam\nnight emergency command','Intrusion warning;\nnight emergency; search','DJI Speaker+Beacon'],
],[20,35,25,20])
sp()

h1('Chapter 6: Unmanned Surface Vessel (Zhigan Times)')
pa('Zhigan Times 3-in-1 USV (mowing/harvesting/detection) as supplementary option for river/canal ports.')
sp()
mt([
    ['Item','Spec','Scenario','Value'],
    ['USV Platform (ZG-WS)','3.2m, 6h battery,\n3m/s max speed, IP65','River/canal patrol;\nwater area maintenance','Replace manned boats;\n80% labor reduction'],
    ['Mowing/Harvest Module','Auto mow + harvest;\n200L capacity;\n0.5-2m depth','Water weed clearing;\nfloat debris collection','10x cleaning efficiency;\n-3 manned boats'],
    ['Water Quality Module','Multi-spectral: COD/turbidity\n/chlorophyll, real-time','Water monitoring;\noil tracking; compliance','20x monitoring efficiency'],
    ['AI Vision Module','Visible + IR dual;\nTaojing AI algorithms','Surface anomaly;\nhazmat water monitor;\noil recognition','Completes air-water\ntotal coverage'],
],[20,35,25,20])
sp()

h1('Chapter 7: Smart Gate System')
pa('Integrated smart gate system for automated vehicle access control with OCR, face ID, approval workflows, and AI linkage.')
sp()
mt([
    ['Module','Technology','Scenario','Value'],
    ['License Plate OCR','Multi-country plate recognition;\nnight IR compensation','Entry/exit gates','<=25s passage;\n>=90% unmanned'],
    ['Container OCR','Box number recognition\n>=99.7% + TOS sync','Gate channels','Auto box verification;\nzero manual errors'],
    ['Driver Face ID','Face + ID OCR linked\nto dispatch system','Gate passage\nhazmat entry','Touchless passage;\n100% unauthorized block'],
    ['Auto Approval','T-Join approval engine;\nmulti-level approval;\nWeChat/email notification','External vehicle access\nhazmat approval\nnight release','30min -> 3min;\ntraceable records'],
    ['Vehicle Damage Record','AI CNN damage detection;\nauto photo on entry/exit','Vehicle handover','Dispute reduction;\nclear liability'],
    ['Gate Health Monitor','IR gate status monitoring;\nanti-crush sensor;\nfault self-diagnosis','All gate points','80% fault reduction;\nzero crush incidents'],
],[20,35,25,20])
sp()
doc.add_page_break()

# ============ CHAPTER 8: SYNERGY ============
h1('Chapter 8: Multi-Capability Synergy & Integration')
pa('Taojing Data Tech T-Join platform serves as the core integration hub, enabling unified multi-vendor device management and closed-loop operations.')
sp()

h2('8.1 T-Join Platform Architecture')
mt([
    ['Layer','Connected Systems','T-Join Function','Value'],
    ['AI Vision','Fixed CCTV (Hikvision/Dahua)\nTaojing 50+ algorithms','Unified algorithm scheduling\nmulti-stream parallel processing\nedge-cloud inference','Single-point alert\ntriggers full-domain response'],
    ['Robot','Qisheng robots\n(EX/AMP/GEN)','Task dispatch engine\nmulti-robot path planning\nauto charge/swap scheduling','Multi-robot collaboration\nblind-spot-free coverage'],
    ['UAV','DJI drones\nexpandable payloads','Route auto-planning\nmulti-sensor data fusion\nauto takeoff/land scheduling','Aerial coverage supplement\nemergency response <=5min'],
    ['USV (Supp.)','Zhigan USV\n(on-demand)','Water task planning\nUAV/robot data fusion\ncompliance report generation','Air-water-land integration\nriver port complete solution'],
    ['Gate','Smart gate system\nplate/box OCR','Vehicle access approval engine\nTOS integration\ngate anomaly linked alert','First defense layer\nlinked with inspection system'],
],[18,25,30,27])
sp()

h2('8.2 End-to-End Closed-Loop Process')
mt([
    ['Step','Unit','T-Join Support','SLA'],
    ['1. Anomaly Detection','Fixed CCTV / Qisheng\n/ DJI UAV','AI auto-detection;\nedge + cloud verification','Real-time (<=100ms)'],
    ['2. Multi-Level Alert','T-Join Platform','Tiered alerting: critical/\nimportant/general;\nWeChat/email/siren','Alert delay <=2s'],
    ['3. Task Dispatch','T-Join Engine','Auto-dispatch to nearest\nrobot/UAV; manual override','Dispatch delay <=10s'],
    ['4. On-Site Action','Qisheng Robot / DJI UAV\n/ Zhigan USV','Result photo upload;\nauto-trigger backup if needed','Response per severity'],
    ['5. Review & Archive','T-Join AI Auto-Review','Before/after AI comparison;\nevent archive; audit trail','Review delay <=60s'],
],[18,25,30,27])
sp()

h2('8.3 Indispensability Proof')
bl('Without Taojing AI: all video requires manual review, miss rate >=95%, no core value',0)
bl('Without Qisheng Robots: hazardous/reefer zones cannot be inspected, extreme human risk',0)
bl('Without DJI UAV: crane tops and large perimeter cannot be quickly inspected, miss rate >=70%',0)
bl('Without Zhigan USV: river/canal water zones uncovered, incomplete solution (supplementary)',0)
bl('Without Smart Gate: vehicle access not automated, approval bottleneck, TOS integration gap',0)
bl('Without T-Join Platform: sub-systems isolated, no synergy effect, greatly reduced overall efficiency',0)
sp()
doc.add_page_break()

# ============ CHAPTER 9: VALUE & ROADMAP ============
h1('Chapter 9: Solution Value & Implementation Roadmap')
pa('This chapter presents the business case, KPI commitments, investment estimates, and phased implementation plan.')
sp()

h2('9.1 Post-Implementation KPI Commitments')
mt([
    ['KPI Category','Metric','Target Value','Current Baseline','Improvement'],
    ['Throughput','Annual TEU capacity increase','+20-30%','100M TEU','High revenue growth'],
    ['Efficiency','Gate passage time','<=25 seconds','60-90 seconds','-60% wait time'],
    ['Safety','Safety incidents','-80%','Baseline year','Dramatic improvement'],
    ['Labor','On-site workforce reduction','-1000 persons','~3000 staff','-33% headcount'],
    ['Inspection','Full-zone inspection coverage','>=99%','~40% coverage','2.5x coverage'],
    ['Response','Emergency response time','<=5 minutes','30+ minutes','-83% response time'],
    ['Uptime','System availability','>=99.9%','~95%','Near-zero downtime'],
    ['Detection','AI detection accuracy','>=99%','~70% (manual)','+29 points'],
],[18,22,22,20,18])
sp()

h2('9.2 Investment Estimate (Baseline: 10M TEU/year)')
mt([
    ['System / Equipment','Vendor','Quantity','Investment (CNY B)'],
    ['Taojing AI Vision System\n(T-Join + 50+ algorithms)','Taojing Data Tech','1 suite','0.8-1.2'],
    ['Qisheng Inspection Robots\n(EX + AMP + GEN)','Hangzhou Qisheng','8-12 units\n+ 3-5 charge stations','0.6-1.0'],
    ['DJI Drone System\n(Dock 2 + expandable payloads)','DJI Innovation','1-2 units\n+ 1 station','0.3-0.5'],
    ['Zhigan USV\n(supplementary, on-demand)','Zhigan Times','1-2 vessels\n(river ports only)','0.1-0.2 (optional)'],
    ['Smart Gate System\n(plate/box OCR + approval)','Hikvision/Dahua','10-15 lanes','0.2-0.4'],
    ['Edge Computing\n(NVIDIA Jetson series)','NVIDIA + Taojing','12-20 units','0.3-0.6'],
    ['TOTAL (excl. Zhigan USV & civil works)','---','---','2.3-3.9'],
],[28,22,22,28])
sp()

h2('9.3 Annual ROI & Payback Period')
mt([
    ['Item','Annual Value (CNY)','Basis'],
    ['Throughput-driven revenue growth','2.5-4.0 billion','+20-30% TEU capacity'],
    ['Labor cost savings','1.2-1.8 billion','-1000 staff @ 120-180K/person'],
    ['Maintenance cost reduction','30-50 million','Predictive maintenance -25~35%'],
    ['AI inspection cost savings','80-120 million','Replace high-risk manual inspection -70%'],
    ['TOTAL ANNUAL BENEFIT','4.6-6.2 billion','---'],
    ['STATIC PAYBACK PERIOD','0.5-0.8 years (6-10 months)','Premium ROI;\nwith depreciation: 1-1.5 years'],
],[35,30,35])
sp()

h2('9.4 Implementation Roadmap (4 Phases)')
mt([
    ['Phase','Timeline','Scope','Investment (CNY B)'],
    ['Phase 1:\nFoundation','Year 1-2','T-Join platform deploy;\nSmart gate install & integrate;\nEdge computing nodes (AI pilot)','0.8-1.5'],
    ['Phase 2:\nRobot & Gate\nFull Rollout','Year 2-3','Qisheng robots deploy;\nGate system full coverage;\nTaojing AI algorithms full load','0.6-1.0'],
    ['Phase 3:\nUAV & Synergy','Year 3-4','DJI drone system online;\nT-Join synergy optimization;\nZhigan USV (on-demand)','0.5-0.8'],
    ['Phase 4:\nOptimization','Year 4+','AI model continuous improvement;\nnew sensor/robot/UAV integration;\nO&M and upgrades','0.2-0.3/year'],
],[15,18,45,22])
sp()

tip([
    'Near-term (6 months): Sign strategic partnership with Taojing, Qisheng, DJI to lock supply chain',
    'Mid-term (1-2 years): Secure 1-2 SE Asia port pilots; prioritize gate + AI vision + robot foundation',
    'Long-term (3-5 years): Establish Smart Port Alliance (Keppel + Taojing + Qisheng + DJI); export to 50+ ports globally',
])
sp()
doc.add_page_break()

# ============ CHAPTER 10: RISKS & MITIGATION ============
h1('Chapter 10: Risks & Mitigation Strategies')
pa('Comprehensive risk analysis with specific mitigation measures across technical, supply chain, implementation, and security dimensions.')
sp()

mt([
    ['Risk Category','Risk Description','Probability','Impact','Mitigation Strategy'],
    ['Technical:\nAI Accuracy','AI algorithm accuracy\nbelow target in complex\nport environments','Medium','High','Phased deployment with\npilot testing;\ncontinuous model training;\nfeedback loop optimization'],
    ['Technical:\nIntegration','Multi-vendor device\nintegration complexity;\nprotocol incompatibility','Medium','High','T-Join standardized API;\npre-integration testing lab;\nvendor SLA enforcement'],
    ['Technical:\nEdge Computing','Edge node overload\nin peak periods;\nlatency spikes','Low','Medium','NVIDIA AGX Orin over-provision;\ncloud fallback mechanism;\nload balancing algorithm'],
    ['Supply Chain:\nComponent','Key component\nshortage (chips, sensors)','Medium','High','6-month forward ordering;\nmulti-source strategy;\nsafety stock policy'],
    ['Supply Chain:\nVendor','Single vendor dependency\nrisk','Medium','Medium','Multi-vendor design;\nopen API standards;\nvendor escrow agreements'],
    ['Implementation:\nTimeline','Project delays due to\nport operational constraints','Medium','Medium','Phased deployment;\nminimal disruption design;\nflexible scheduling'],
    ['Implementation:\nUser Adoption','Port staff resistance\nto new technology','Low','Medium','Comprehensive training program;\ngradual transition;\nincentive alignment'],
    ['Security:\nCyber','Cyber attacks on port\nOT/IT infrastructure','Low','Critical','Network segmentation;\nend-to-end encryption;\npenetration testing;\nincident response plan'],
    ['Security:\nPhysical','Equipment damage\nfrom harsh port environment','Medium','Medium','IP65/67 rated equipment;\nreplaceable modules;\nremote diagnostics'],
    ['Regulatory:\nCompliance','Port regulations vary\nacross SE Asia countries','Medium','Medium','Local compliance assessment\nper country;\nlegal advisory retained;\nregulatory liaison'],
],[15,20,12,12,41])
sp()
doc.add_page_break()

# ============ CHAPTER 11: AFTER-SALES & O&M ============
h1('Chapter 11: After-Sales Service & O&M Commitments')
pa('Comprehensive service commitments ensuring long-term system reliability and performance.')
sp()

h2('11.1 Warranty & SLA')
mt([
    ['Service Item','Standard','Premium','Response Time'],
    ['Hardware Warranty','3 years','5 years','---'],
    ['Software Maintenance','3 years included','5 years included','---'],
    ['On-Site Support','24x7 response','24x7 response\n+ dedicated engineer','<=4 hours on-site'],
    ['Remote Support','24x7 hotline','24x7 priority hotline','<=30 minutes response'],
    ['Emergency Repair','Within 24 hours','Within 12 hours\n+ backup equipment','Critical: <=4h'],
    ['System Availability SLA','>=99.5%','>=99.9%','---'],
    ['AI Model Update','Quarterly','Monthly','---'],
    ['Spare Parts Supply','72 hours local stock','48 hours local stock\n+ regional hub','---'],
],[20,25,30,25])
sp()

h2('11.2 Preventive Maintenance Schedule')
mt([
    ['Equipment','Inspection Freq.','Maintenance Scope','Responsible Party'],
    ['T-Join Platform','Monthly','Software update, security patch,\nperformance tuning','Taojing Data Tech'],
    ['AI Algorithms','Quarterly','Model retraining, accuracy audit,\nnew algorithm deployment','Taojing Data Tech'],
    ['Qisheng Robots','Weekly visual,\nMonthly deep','Mechanical check, battery health,\nsensor calibration','Hangzhou Qisheng\n+ local team'],
    ['DJI Drones','Before each flight\n+ monthly deep','Airframe integrity, payload calibration,\nbattery cycle management','DJI + local team'],
    ['Smart Gates','Daily auto-check,\nMonthly manual','Gate mechanism, camera cleaning,\nOCR accuracy verification','Local team\n+ vendor support'],
    ['Edge Computing','Monthly','Hardware health, cooling system,\nstorage optimization','Taojing + NVIDIA'],
    ['Network Infrastructure','Monthly','5G/WiFi signal coverage,\nbandwidth optimization','Local ISP + Taojing'],
],[20,20,30,30])
sp()

h2('11.3 Training Plan')
mt([
    ['Training Type','Target Audience','Duration','Content'],
    ['System Overview','Port management','1 day','Architecture, capabilities,\nKPI dashboards'],
    ['Operator Training','Day-to-day operators','3 days','T-Join console, alert handling,\ntask management, basic troubleshooting'],
    ['Maintenance Training','Technical staff','5 days','Equipment maintenance, sensor\ncalibration, spare part replacement'],
    ['Advanced Training','IT/engineering team','10 days','System admin, API integration,\nAI model management, security'],
    ['Emergency Response','All operators + mgmt','1 day','Emergency procedures, escalation\npaths, communication protocols'],
    ['Refresher Training','All staff','Bi-annual (1 day)','New features, best practices,\nlessons learned review'],
],[20,20,20,40])
sp()
doc.add_page_break()

# ============ APPENDIX A: GLOSSARY ============
h1('Appendix A: Glossary & Abbreviations')
mt([
    ['Abbreviation','Full Term','Definition'],
    ['AGV','Automated Guided Vehicle','Autonomous vehicle for container transfer in automated terminals'],
    ['ASC','Automated Stacking Crane','Crane for automated container stacking/retrieval in yard blocks'],
    ['RMQC','Remote-Manned Quay Crane','Ship-to-shore crane operated remotely from control center'],
    ['TOS','Terminal Operating System','Core software managing all terminal operations and planning'],
    ['AI','Artificial Intelligence','Computer systems performing tasks requiring human intelligence'],
    ['IR','Infrared','Electromagnetic radiation beyond visible spectrum for thermal imaging'],
    ['OCR','Optical Character Recognition','Technology for converting images of text into machine-readable data'],
    ['5G','5th Generation Mobile Network','High-speed, low-latency wireless network for IoT and real-time data'],
    ['IoT','Internet of Things','Network of connected sensors and devices for data collection'],
    ['Digital Twin','Digital Twin','Virtual 3D replica of physical assets for simulation and optimization'],
    ['USV','Unmanned Surface Vessel','Autonomous or remotely-controlled watercraft for surface operations'],
    ['UAV','Unmanned Aerial Vehicle','Drone or autonomous aircraft for aerial operations'],
    ['SLA','Service Level Agreement','Contractual commitment to service performance metrics'],
    ['TEU','Twenty-foot Equivalent Unit','Standard measure of container capacity (one 20ft container)'],
    ['Edge Computing','Edge Computing','Processing data near the source rather than in centralized cloud'],
    ['NVIDIA Jetson','NVIDIA Jetson','Edge AI computing platform for embedded inference'],
    ['PSA','Port of Singapore Authority','Singapore\'s port operator, managing Tuas Port'],
    ['SGX','Singapore Exchange','Singapore\'s securities exchange where Keppel is listed'],
    ['CV','Computer Vision','AI field enabling machines to interpret visual information from cameras'],
    ['LiDAR','Light Detection and Ranging','Sensor using laser pulses for 3D mapping and distance measurement'],
    ['KPI','Key Performance Indicator','Measurable value demonstrating effectiveness of an activity'],
    ['ROI','Return on Investment','Financial metric comparing investment cost to generated returns'],
    ['Capex','Capital Expenditure','Funds used to acquire, upgrade, or maintain physical assets'],
    ['OT','Operational Technology','Hardware/software for industrial process monitoring and control'],
    ['IT','Information Technology','Systems for data processing, storage, and communication'],
    ['BMS','Building Management System','Centralized control system for building facilities and equipment'],
],[18,30,52])
sp()

# ============ BACK COVER ============
sp();sp();sp()
cover_box([
    ('Thank You',24,"AACCEE",True,False),
    (' ',8,"AACCEE",False,False),
    ('We look forward to partnering with Keppel Corporation',12,"AACCEE",False,False),
    ('to build the next generation of smart ports across Southeast Asia.',12,"AACCEE",False,False),
    (' ',12,"AACCEE",False,False),
    ('For inquiries, please contact:',10,DB,False,False),
    ('Keppel Corporation - Business Development Team',10,MB,False,False),
    (' ',8,"AACCEE",False,False),
    ('Keppel Corporation - Creating solutions for a sustainable future',9,"666666",False,True),
], LB, MB, 350)

# ============ FOOTER ============
add_footer()

# ============ SAVE ============
out = '/Users/tom/.qclaw/workspace/Smart_Port_AI_Solution_Keppel_V5_Complete.docx'
doc.save(out)
print('OK: ' + out)

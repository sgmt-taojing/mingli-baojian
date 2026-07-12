# -*- coding: utf-8 -*-
# 智慧港口AI巡检解决方案 V5 - 去AI味版
# 优化方向：去掉emoji、营销话术；句式多样化；用词克制；像人写的方案
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.page_width  = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin= Cm(2.5); s.right_margin = Cm(2.5)

# ── 工具函数 ──────────────────────────────────────────────────────────────
def hc(h): return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))
DB,CG,ORG,LG,DG,BG,DKBG,MB,LR,LB = "000000","00A896","E76F51","F2F2F2","1E293B","EBF5FB","0D2F4F","028090","AACCEE","F8F9FA"

def bl(t, sz=10, sp=2):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after=Pt(sp)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(sz); r.font.color.rgb=hc(DB)

def pa(t, sz=9, sp=2):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after=Pt(sp)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(sz); r.font.color.rgb=hc("444444")

def h1(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(13); r.bold=True; r.font.color.rgb=hc(MB)

def h2(t):
    p=doc.add_paragraph(t); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(11); r.bold=True; r.font.color.rgb=hc(DB)

def sp(): doc.add_paragraph()

def bc2(t,c):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),c)
    tcPr.append(shd)

def cm(t,top,right,bottom,left):
    tc=t._tc; tcPr=tc.get_or_add_tcPr()
    tcMar=OxmlElement('w:tcMar')
    for side,val in [('top',top),('right',right),('bottom',bottom),('left',left)]:
        m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),str(int(val*20))); m.set(qn('w:type'),'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def hdr_cell(cell, text, fill=MB, fc="FFFFFF", fs=8.5):
    p=cell.paragraphs[0]; p.text=text; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(fs); r.bold=True; r.font.color.rgb=hc(fc)
    bc2(cell, fill)

def dat_cell(cell, text, fs=8, color=DB, bold=False):
    p=cell.paragraphs[0]; p.text=text; p.paragraph_format.space_after=Pt(1)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(fs); r.font.color.rgb=hc(color)
    if bold: p.paragraphs[0].runs[0].bold=True

def add_footer():
    for s in doc.sections:
        f=s.footer; f.is_linked_to_previous=False
        p=f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run('吉宝集团（Keppel Corporation）| 机密文件 | 2026年5月')
        r.font.name='SimSun'; r.font.size=Pt(7); r.font.color.rgb=hc("999999")

# ── 封面 ──────────────────────────────────────────────────────────────────
tc0 = doc.add_table(rows=1,cols=1).cell(0,0)
bc2(tc0, DKBG); cm(tc0,400,400,400,400)
for line, sz, col, bold in [
    ('智慧港口AI巡检预警解决方案', 20, LR, True),
    (' ', 6, LR, False),
    ('技术建议书  V5.0', 15, "AACCEE", True),
    (' ', 10, LR, False),
    ('牵头方：淘景数科（总包/方案整合）  配套供应商：大疆创新 · 杭州旗晟 · 智感时代', 9, "AACCEE", False),
    (' ', 6, LR, False),
    ('为吉宝集团（Keppel）定制', 12, DB, True),
    (' ', 4, LR, False),
    ('2026年5月  |  版本 V5.0', 9, "666666", False),
    ('内部资料，请勿外传', 10, ORG, True),
]:
    p=tc0.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='Arial' if col==LR else 'SimSun'; r.font.size=Pt(sz); r.font.color.rgb=hc(col)
    if line.startswith(' '): p.paragraph_format.space_after=Pt(1)
    else: p.paragraph_format.space_after=Pt(4)
sp()
tc1 = doc.add_table(rows=1,cols=1).cell(0,0)
bc2(tc1, LB); cm(tc1,250,250,250,250)
for line in [
    '淘景数科负责总体方案设计、AI算法及T-Join平台整合交付',
    '大疆创新、杭州旗晟、智感时代为指定配套设备供应商',
    '交付形式可根据吉宝需求选择整体交付或按模块分步实施',
]:
    p=tc1.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(9); r.font.color.rgb=hc(MB)
    p.paragraph_format.space_after=Pt(2)
doc.add_page_break()

# ── 目录 ──────────────────────────────────────────────────────────────────
h1('目录（页面卷标，点击跳转）')
pa('本方案共15章，以淘景数科为总包方，整合多载体AI巡检能力，为吉宝集团提供可落地的港口智能化方案。')
sp()
toc = doc.add_table(rows=15, cols=2); toc.style='Table Grid'; toc.alignment=WD_TABLE_ALIGNMENT.CENTER
chapters = [
    ('第0章','吉宝企业调研与东南亚市场机会'),
    ('第1章','行业标杆案例（14个国内外案例）'),
    ('第2章','总包定位与供应商分工'),
    ('第3章','全场景AI视觉算法矩阵'),
    ('第4章','地面巡检机器人（杭州旗晟）'),
    ('第5章','空中无人机巡检（大疆创新）'),
    ('第6章','无人船补充方案（智感时代）'),
    ('第7章','港口道闸设备（智能通行管理）'),
    ('第8章','多载体协同体系（T-Join平台）'),
    ('第9章','方案价值、KPI承诺与实施路径'),
    ('第10章','风险分析与对策'),
    ('第11章','售后服务与运维承诺'),
    ('附录A','术语表与缩写对照'),
]
for i,(ch,title) in enumerate(chapters):
    cr=toc.rows[i].cells
    p=cr[0].paragraphs[0]; p.text=ch; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='Microsoft YaHei'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc(MB)
    p=cr[1].paragraphs[0]; p.text=title; p.paragraph_format.space_after=Pt(2)
    for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(9); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ── 第0章 ────────────────────────────────────────────────────────────────
h1('第0章：吉宝企业调研与背景分析')
h2('0.1 吉宝集团业务构成')
tbl = doc.add_table(rows=4, cols=3); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['业务领域','主要方向','与港口的关联']):
    hdr_cell(tbl.rows[0].cells[i], h, MB)
rows = [
    ['基础设施（Infrastructure）','港口投资与运营、能源与数据中心、资产管理','大士港投资运营；东南亚港口网络'],
    ['房地产（Real Estate）','商业地产开发、城市综合体','港口配套园区、客户服务中心'],
    ['连接性（Connectivity）','ICT与网络基础设施、智慧城市','港口5G专网、数据中心、数字化服务'],
]
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        dat_cell(tbl.rows[ri+1].cells[ci], val)
h2('0.2 大士港（Tuas Port）— 全球最大全自动化码头')
bl('总投资35亿新元，预计2030年全面完工，设计吞吐能力6500万TEU/年')
bl('采用AGV自动导引车 + ASC自动堆垛起重机 + 远程岸桥，全流程无人化操作')
bl('AI智能调度系统覆盖船舶到港预测、堆场自动分配、设备协同调度三个层面')
bl('已建成数字孪生系统，全港区三维建模，实时反映物理码头运行状态')
bl('吉宝是大士港主要投资方与运营方之一，对自动化码头技术有实际运营经验')
h2('0.3 吉宝全球港口布局')
bl('新加坡：大士港（核心项目）、PSA港口（合作运营）')
bl('中国：上海、宁波、广州、天津等港口均有投资或合作项目')
bl('东南亚：越南（海防港、胡志明港）、泰国（林查班港）、马来西亚（巴生港）、印尼（雅加达港）、菲律宾（马尼拉港）')
bl('全球影响力：吉宝港口网络覆盖约20个港口，在东南亚地区具备较强的政商资源')
h2('0.4 东南亚市场机会')
seatbl = doc.add_table(rows=7, cols=5); seatbl.style='Table Grid'; seatbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['国家','重点港口','可落地场景','吉宝优势','建议优先级']):
    hdr_cell(seatbl.rows[0].cells[i], h, MB)
sea = [
    ['越南','海防港、胡志明港','岸桥监控、堆场管理、航道监测','吉宝在越有港口投资','高'],
    ['泰国','林查班港、曼谷港','无人机巡逻、危化品检测','吉宝泰国分公司','高'],
    ['马来西亚','巴生港、丹戎帕拉帕斯港','全场景AI（对标大士港）','吉宝马来西亚分公司','中'],
    ['印尼','雅加达港、泗水港','无人机+机器人（群岛地形适合）','吉宝印尼分公司','中'],
    ['菲律宾','马尼拉港、宿务港','基础AI（航道监测、闸口OCR）','吉宝菲律宾分公司','低'],
    ['新加坡','大士港、PSA港口','全自动化码头升级、数字孪生','吉宝总部所在地','高（示范）'],
]
for ri, row in enumerate(sea):
    for ci, val in enumerate(row):
        dat_cell(seatbl.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第1章 ────────────────────────────────────────────────────────────────
h1('第1章：行业标杆案例')
h2('1.1 国内案例')
tbl1 = doc.add_table(rows=10, cols=4); tbl1.style='Table Grid'; tbl1.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['港口/项目','关键指标','AI/自动化供应商','借鉴要点']):
    hdr_cell(tbl1.rows[0].cells[i], h, MB)
cases = [
    ['青岛港全自动化码头','作业效率62.62箱/小时（世界纪录）','海康威视、海大宇航','自研系统 + 全域AI覆盖'],
    ['上海洋山港四期','年吞吐745万TEU，全球最大单体自动化码头','华为、振华重工','全局调度算法'],
    ['宁波舟山港','AI+机器狗空箱查验，耗时从2小时降至10分钟','海康机器人','机器狗替代人工高危作业'],
    ['南京港','AI平台，AI识别准确率95%+','华为、淘景数科','统一管理平台架构'],
    ['山东港口天和','设备综合效率提升30%，等待时间减少90%','华为、易核验','轻量化改造方案'],
    ['烟台港','干散货全自动化码头，综合效率提升8%','华为、北斗','全流程无人化'],
    ['镇江港','AI皮带检测，年节省运维成本300万以上','海康威视','预测性维护'],
    ['唐山港','四足消防机器人+80L/s水炮，复杂环境巡检','宇树科技、海康威视','消防+巡检融合'],
    ['舟山甬舟码头','空地协同，机器人+无人机联合巡检','海康机器人','多载体联动'],
]
for ri, row in enumerate(cases):
    for ci, val in enumerate(row):
        dat_cell(tbl1.rows[ri+1].cells[ci], val)
h2('1.2 国际案例')
tbl2 = doc.add_table(rows=5, cols=4); tbl2.style='Table Grid'; tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['港口/项目','关键指标','技术特色','借鉴要点']):
    hdr_cell(tbl2.rows[0].cells[i], h, CG)
intl = [
    ['新加坡大士港','设计6500万TEU/年，全球最大','数字孪生+全无人化','吉宝已参与，技术方案可复用'],
    ['荷兰鹿特丹港','欧洲最大自动化港口','AI调度+区块链','欧洲标准，适合高端客户'],
    ['韩国釜山港','年吞吐800万TEU，全球第7','智能闸口+AI船舶预测','东亚地区竞争标杆'],
    ['埃及苏赫纳港','北非首座全自动集装箱码头','振华重工+中控技术','中国方案出海成功案例'],
]
for ri, row in enumerate(intl):
    for ci, val in enumerate(row):
        dat_cell(tbl2.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第2章 ────────────────────────────────────────────────────────────────
h1('第2章：总包定位与供应商分工')
h2('2.1 总包方：淘景数科')
bl('淘景数科作为本方案的总包方，承担以下职责：')
bl('  （1）AI算法提供：覆盖固定摄像机、无人机、机器人、无人船四大载体的50+场景算法；')
bl('  （2）T-Join整合平台：实现多厂商设备的统一接入、任务调度和数据管理；')
bl('  （3）总体项目管理：协调各配套供应商，确保交付进度和质量；')
bl('  （4）售后服务总负责：作为单一责任主体，处理吉宝的所有运维需求。')
bl('交付形式灵活：可选择整体交付（一站式），也可按模块分步实施（可拆卸），淘景负责接口标准统一。')
h2('2.2 配套供应商及分工')
ptbl = doc.add_table(rows=6, cols=3); ptbl.style='Table Grid'; ptbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['角色','公司','主要贡献']):
    hdr_cell(ptbl.rows[0].cells[i], h, MB)
pdata = [
    ['总包/整合方','淘景数科','AI算法（50+场景）+ T-Join平台 + 总体交付管理'],
    ['配套供应商（无人机）','大疆创新','Matrice 350 RTK工业无人机 + Dock 2自动基站 + 多传感器载荷'],
    ['配套供应商（机器人）','杭州旗晟','防爆型/水陆两栖/通用型巡检机器人'],
    ['配套供应商（无人船）','智感时代','割草/打捞/检测一体无人船（适用于内河及港口水域）'],
    ['配套供应商（道闸）','智能道闸厂商','车牌/箱号/人脸OCR + 自动审批引擎'],
]
for ri, row in enumerate(pdata):
    for ci, val in enumerate(row):
        dat_cell(ptbl.rows[ri+1].cells[ci], val)
h2('2.3 选择淘景数科作为总包的理由')
bl('（1）算法覆盖全面：50+港口场景算法，四大载体均有成熟方案，不需要吉宝分别对接多家算法厂商；')
bl('（2）整合经验丰富：T-Join平台已对接20+品牌设备，真正实现一个平台管理所有硬件；')
bl('（3）国内案例可查：青岛港、宁波舟山港等14个标杆项目均有公开信息可以核实；')
bl('（4）技术路线清晰：边缘计算（NVIDIA Jetson）+ 云端训练，模型可持续优化；')
bl('（5）售后服务网络：全国主要港口城市均有服务团队，重大故障承诺4小时到场。')
doc.add_page_break()

# ── 第3章 ────────────────────────────────────────────────────────────────
h1('第3章：全场景AI视觉算法矩阵')
h2('3.1 固定摄像机算法（7大分区，22+算法）')
tbl31 = doc.add_table(rows=9, cols=4); tbl31.style='Table Grid'; tbl31.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['分区','代表算法','准确率','联动方式']):
    hdr_cell(tbl31.rows[0].cells[i], h, MB)
algo_fix = [
    ['岸线（岸桥）','箱号OCR识别、吊具对位、人员入侵检测','≥99.5%','箱号不匹配→道闸自动拦截'],
    ['堆场','AI盘存、堆垛稳定性分析、火灾预警','≥99%','火情确认→无人机携带灭火弹处置'],
    ['闸口','车牌OCR、箱号校验、人脸核验','≥99.5%','未授权车辆→道闸锁死并记录'],
    ['危化品区','防护装备检测、气体泄漏检测（可见光+红外+超声波）','≥97%','泄漏预警→防爆机器人近距离核查'],
    ['周界','入侵检测、围栏损坏检测','≥99%','入侵事件→机器人+无人机联合处置'],
    ['配电室/机房','设备过热检测（红外热成像）','≥98%','超温预警→机器人现场核查'],
    ['航道','油污检测、船舶异常行为监测','≥97%','油污事件→无人船采样+无人机追踪'],
    ['全港区通用','反光衣/安全帽/吸烟/打电话检测','≥99%','违规事件→人工复核'],
]
for ri, row in enumerate(algo_fix):
    for ci, val in enumerate(row):
        dat_cell(tbl31.rows[ri+1].cells[ci], val)
h2('3.2 无人机算法（7大分区，16+算法）')
tbl32 = doc.add_table(rows=9, cols=4); tbl32.style='Table Grid'; tbl32.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['分区','代表算法','准确率','联动方式']):
    hdr_cell(tbl32.rows[0].cells[i], h, CG)
algo_uav = [
    ['岸桥顶部','设备巡检（钢丝绳/滑轮磨损检测）','≥98%','磨损超限→自动生成维护工单'],
    ['堆场','红外全景扫描、堆位自动盘存','≥99%','热点发现→两栖机器人核查'],
    ['周界','全景巡逻（覆盖固定摄像机盲区）','≥99%','入侵确认→喊话威慑+机器人处置'],
    ['危化品区','360°风险评估（多传感器融合）','≥97%','高风险确认→防爆机器人处置'],
    ['航道','多光谱巡逻（油污/藻华检测）','≥97%','油污发现→无人船采样确认'],
    ['应急处置','快速侦察（载荷可切换）','≥99%','重大事故→多载体协同调度'],
    ['设备清洗','岸桥/摄像机自动清洗','≥90%','按计划自动执行清洗航线'],
    ['数据回传','边缘计算节点，实时回传','≥99.9%','所有数据统一汇入T-Join平台'],
]
for ri, row in enumerate(algo_uav):
    for ci, val in enumerate(row):
        dat_cell(tbl32.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第4-6章（精简，避免篇幅过长）────────────────────────────────────
h1('第4-6章：配套设备能力说明')
h2('4.1 地面巡检机器人（杭州旗晟，配套供应商）')
rtbl = doc.add_table(rows=5, cols=4); rtbl.style='Table Grid'; rtbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['型号','适用区域','核心能力','续航']):
    hdr_cell(rtbl.rows[0].cells[i], h, MB)
robots = [
    ['QS-EX（防爆型）','危化品区、油料区','红外+气体+可见光三重融合检测','8小时（支持自动换电）'],
    ['QS-AMP（水陆两栖）','冷藏箱区、电缆隧道','插头接触检测+电缆损伤识别','10小时'],
    ['QS-GEN（通用型）','堆场、周界、配电室','火灾扫描+入侵巡逻','12小时'],
    ['自动充电坞','—','对接精度±5mm，全自动换电','—'],
]
for ri, row in enumerate(robots):
    for ci, val in enumerate(row):
        dat_cell(rtbl.rows[ri+1].cells[ci], val)
h2('5.1 工业无人机（大疆创新，配套供应商）')
bl('主设备：Matrice 350 RTK，IP55防护，抗风7级，续航55分钟，支持九向避障。')
bl('自动基站：Dock 2，IP55防护，支持自动充电、自动换电、气象监测。')
bl('可扩展载荷：灭火弹发射器、水炮、多光谱传感器、空气质量传感器、水质传感器。')
h2('6.1 无人船（智感时代，配套供应商，按需选配）')
bl('适用于内河港口或水域面积较大的港口（新加坡大士港水域情况复杂，有较大应用空间）。')
bl('核心功能：水面割草、垃圾打捞、水质检测三合一，替代人工船，人力成本降低约80%。')
doc.add_page_break()

# ── 第7章 道闸 ───────────────────────────────────────────────────────────
h1('第7章：港口道闸设备（智能通行管理）')
h2('7.1 系统功能模块')
gtbl = doc.add_table(rows=8, cols=3); gtbl.style='Table Grid'; gtbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['功能模块','说明','准确率/性能']):
    hdr_cell(gtbl.rows[0].cells[i], h, ORG)
gdata = [
    ['车牌OCR识别','支持中新马等多国车牌，与TOS系统联动','≥99.5%'],
    ['箱号OCR校验','与TOS系统同步校验，箱号不匹配自动拦截','≥99.7%'],
    ['人脸核验','与派车系统联动，未授权驾驶员禁止入港','≥99.9%'],
    ['自动审批引擎','多级审批流程，微信/邮件通知，审批时效从30分钟缩短至3分钟','≥99%'],
    ['车辆损伤记录','入场/出场AI损伤检测对比，减少纠纷','≥95%'],
    ['道闸健康监控','红外状态监测+防砸传感器，故障率降低约80%','≥99%'],
    ['与T-Join平台联动','AI事件自动触发道闸动作，形成闭环','≥99%'],
]
for ri, row in enumerate(gdata):
    for ci, val in enumerate(row):
        dat_cell(gtbl.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第8章 协同 ──────────────────────────────────────────────────────────
h1('第8章：多载体协同体系（T-Join平台核心能力）')
h2('8.1 T-Join平台简介')
bl('T-Join平台由淘景数科自主研发，是方案的核心整合中枢，主要功能包括：')
bl('（1）统一接入：已对接20+品牌设备（大疆/旗晟/智感/海康/大华等），同一套平台管理所有硬件；')
bl('（2）协同调度：AI事件触发后，平台自动调度距离最近的载体（无人机/机器人/无人船）到场处置；')
bl('（3）闭环管理：事件检测→任务派发→现场处置→结果回传→人工复核，全程可追溯；')
bl('（4）数据分析：所有事件记录、处置结果、设备状态统一存储，自动生成运维报告。')
h2('8.2 多载体联动场景（T-Join平台统一调度）')
ltbl = doc.add_table(rows=11, cols=4); ltbl.style='Table Grid'; ltbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['联动场景','触发载体','协同载体（T-Join调度）','目标响应时效']):
    hdr_cell(ltbl.rows[0].cells[i], h, "E76F51")
linkage = [
    ['堆场火情','固定摄像机','无人机（灭火弹）+ 机器人（核查）','≤5分钟'],
    ['堆垛倾斜','固定摄像机','机器人（测量）+ 无人机（全景确认）','≤3分钟'],
    ['危化品泄漏','固定摄像机','防爆机器人（近距离）+ 无人机（360°评估）','≤8分钟'],
    ['周界入侵','固定摄像机','机器人（拦截）+ 无人机（追踪）','≤2分钟'],
    ['航道油污','固定摄像机','无人船（采样确认）+ 无人机（扩散追踪）','≤15分钟'],
    ['岸桥设备异常','固定摄像机','无人机（近距离检查）+ 机器人（地面支持）','≤5分钟'],
    ['道闸故障/尾随','固定摄像机','机器人（人工核验）','≤3分钟'],
    ['冷藏箱温度异常','固定摄像机','两栖机器人（插头检测）+ 无人机（红外确认）','≤10分钟'],
    ['全域空气质量监测','固定摄像机（扩展）','无人机（多传感器）+ 无人船（水质）','每日2次'],
]
for ri, row in enumerate(linkage):
    for ci, val in enumerate(row):
        dat_cell(ltbl.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第9章 KPI + 投资 ───────────────────────────────────────────────────
h1('第9章：方案价值、KPI承诺与实施路径')
h2('9.1 KPI承诺（淘景数科作为总包方，验收标准可写入合同）')
ktbl = doc.add_table(rows=10, cols=4); ktbl.style='Table Grid'; ktbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['类别','指标名称','承诺值（验收标准）','测量方式']):
    hdr_cell(ktbl.rows[0].cells[i], h, MB)
kpi = [
    ['吞吐能力','港口吞吐量提升','+20~30%','TOS系统数据对比'],
    ['通行效率','闸口单车通行时间','≤25秒/车','道闸系统日志'],
    ['安全改善','港口安全事故率降低','-80%','安全记录年度对比'],
    ['人力优化','现场作业人员需求减少','-33%','人力成本统计'],
    ['AI性能','AI检测准确率（大部分场景）','≥99%','人工抽检'],
    ['应急响应','从AI检测到现场处置的平均时间','≤5分钟','事件响应日志'],
    ['巡检覆盖','全港区巡检覆盖率','≥99%','巡检报告统计'],
    ['系统稳定','系统可用性（SLA）','≥99.9%','系统监控日志'],
    ['投资回报','静态投资回收期','0.5~0.8年（6~10个月）','财务测算模型'],
]
for ri, row in enumerate(kpi):
    for ci, val in enumerate(row):
        dat_cell(ktbl.rows[ri+1].cells[ci], val)
h2('9.2 投资测算（总包：淘景数科）')
itbl = doc.add_table(rows=8, cols=4); itbl.style='Table Grid'; itbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['系统/设备','供应商','规模','投资估算（亿元）']):
    hdr_cell(itbl.rows[0].cells[i], h, CG)
idata = [
    ['AI视觉系统（T-Join + 50+算法）','淘景数科（总包）','1套','0.8~1.2'],
    ['无人机系统（Matrice 350 + Dock 2）','大疆创新（配套）','3~5台+2基站','0.3~0.5'],
    ['巡检机器人（防爆+两栖+通用）','杭州旗晟（配套）','8~12台','0.4~0.7'],
    ['无人船（割草+打捞+检测）','智感时代（配套）','2~4台','0.15~0.25'],
    ['智能道闸系统','道闸厂商（配套）','10~20套','0.2~0.4'],
    ['实施+培训+3年运维','淘景数科（总包）','1项','0.6~1.0'],
    ['合计','—','—','2.3~3.9'],
]
for ri, row in enumerate(idata):
    for ci, val in enumerate(row):
        c = itbl.rows[ri+1].cells[ci]
        p=c.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
        if ri==6 and ci==3:
            for r in p.runs: r.bold=True; r.font.color.rgb=hc(ORG)
h2('9.3 实施路线图（淘景数科总体协调）')
rtbl2 = doc.add_table(rows=6, cols=4); rtbl2.style='Table Grid'; rtbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['阶段','时间','关键任务（淘景总包协调）','投资']):
    hdr_cell(rtbl2.rows[0].cells[i], h, MB)
roadmap = [
    ['第一阶段：基础建设','第1~2年','T-Join部署；智能道闸安装；AI试点（1-2个分区）','0.8~1.5亿'],
    ['第二阶段：机器人与道闸铺开','第2~3年','旗晟机器人部署；道闸系统全覆盖；AI算法全量加载','0.6~1.0亿'],
    ['第三阶段：无人机与协同优化','第3~4年','大疆无人机上线；T-Join协同优化；智感无人船按需部署','0.5~0.8亿'],
    ['第四阶段：持续优化与扩展','第4年及以后','AI模型持续改进；新传感器集成；东南亚港口扩展','0.2~0.3亿/年'],
]
for ri, row in enumerate(roadmap):
    for ci, val in enumerate(row):
        c = rtbl2.rows[ri+1].cells[ci]
        p=c.paragraphs[0]; p.text=val; p.paragraph_format.space_after=Pt(1)
        for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
doc.add_page_break()

# ── 第10章 风险 ────────────────────────────────────────────────────────
h1('第10章：风险分析与对策')
rstbl = doc.add_table(rows=7, cols=3); rstbl.style='Table Grid'; rstbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['风险类别','具体风险','对策（淘景总包负责协调）']):
    hdr_cell(rstbl.rows[0].cells[i], h, ORG)
risks = [
    ['技术风险','AI算法准确率不达预期','分阶段部署，先试点验证；持续模型训练；KPI不达标免费优化'],
    ['供应链风险','关键零部件短缺（芯片/传感器）','提前6个月下单；多源供应策略；淘景协调安全库存'],
    ['实施风险','项目进度延误','分阶段部署；最小干扰设计；灵活调度资源'],
    ['安全风险','网络攻击、数据泄露','网络分区隔离；端到端加密；定期渗透测试'],
    ['合规风险','东南亚各国监管差异','逐国合规评估；聘请当地法律顾问；建立监管联络机制'],
    ['财务风险','投资超预算','固定总价合同；月度财务报告；严格的变更管理流程'],
]
for ri, row in enumerate(risks):
    for ci, val in enumerate(row):
        dat_cell(rstbl.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 第11章 售后 ────────────────────────────────────────────────────────
h1('第11章：售后服务与运维承诺（淘景数科总包负责）')
h2('11.1 服务等级协议（SLA）')
slatbl = doc.add_table(rows=8, cols=3); slatbl.style='Table Grid'; slatbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['服务项目','标准版SLA','高级版SLA（推荐）']):
    hdr_cell(slatbl.rows[0].cells[i], h, MB)
sla = [
    ['硬件保修（无人机/机器人/道闸）','3年','5年（淘景协调各配套供应商）'],
    ['现场技术支持','7×24响应，72小时到场','7×24响应，4小时到场'],
    ['远程技术支持','7×24热线，≤30分钟响应','7×24优先热线，≤15分钟响应'],
    ['AI算法模型更新与优化','每季度1次','每月1次（含新算法部署）'],
    ['系统可用性承诺','≥99.5%','≥99.9%（高级版）'],
    ['备件供应保障','72小时（本地库存）','48小时（本地+区域中心仓）'],
    ['T-Join平台软件更新','每月安全补丁','每月功能更新+安全补丁'],
]
for ri, row in enumerate(sla):
    for ci, val in enumerate(row):
        dat_cell(slatbl.rows[ri+1].cells[ci], val)
h2('11.2 培训计划（淘景数科负责交付）')
trtbl = doc.add_table(rows=7, cols=3); trtbl.style='Table Grid'; trtbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['培训类型','时长/频率','培训内容']):
    hdr_cell(trtbl.rows[0].cells[i], h, CG)
training = [
    ['系统概述培训','1天（项目启动前）','智慧港口AI理念；方案整体架构；T-Join平台功能介绍'],
    ['操作培训','3天（设备交付时）','固定摄像机查看；无人机/机器人操作；道闸管理'],
    ['维护培训','5天（设备交付后）','AI算法调优；设备故障诊断；T-Join平台日常运维'],
    ['高级培训','10天（高级SLA客户）','Python/AI基础；模型再训练；定制开发入门'],
    ['应急演练','1天（每6个月）','火情应急响应；危化品泄漏处置；周界入侵处置'],
    ['复训','0.5天（每6个月）','新功能介绍；最佳实践分享；常见问题解答'],
]
for ri, row in enumerate(training):
    for ci, val in enumerate(row):
        dat_cell(trtbl.rows[ri+1].cells[ci], val)
doc.add_page_break()

# ── 附录 ─────────────────────────────────────────────────────────────────
h1('附录A：术语表与缩写对照')
terms = [['AGV','自动导引车（Automated Guided Vehicle）'],['ASC','自动堆垛起重机'],['RMQC','远程操控岸桥'],['TOS','码头操作系统'],['AI','人工智能'],['OCR','光学字符识别'],['5G','第五代移动通信'],['IoT','物联网'],['UAV','无人驾驶航空器（无人机）'],['USV','无人水面艇'],['SLA','服务等级协议'],['TEU','二十英尺等效单位'],['ROI','投资回报率'],['NVIDIA Jetson','英伟达边缘AI计算平台']]
terms_per_row = 4; total_rows = (len(terms)+terms_per_row-1)//terms_per_row
termtbl = doc.add_table(rows=total_rows, cols=terms_per_row*2); termtbl.style='Table Grid'; termtbl.alignment=WD_TABLE_ALIGNMENT.CENTER
for ri in range(total_rows):
    for ci in range(terms_per_row):
        idx = ri*terms_per_row + ci
        if idx < len(terms):
            abbr_cell = termtbl.rows[ri].cells[ci*2]
            p=abbr_cell.paragraphs[0]; p.text=terms[idx][0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.name='Arial'; r.font.size=Pt(9); r.bold=True; r.font.color.rgb=hc(MB)
            bc2(abbr_cell, "EBF5FB")
            mean_cell = termtbl.rows[ri].cells[ci*2+1]
            p=mean_cell.paragraphs[0]; p.text=terms[idx][1]; p.paragraph_format.space_after=Pt(1)
            for r in p.runs: r.font.name='SimSun'; r.font.size=Pt(8); r.font.color.rgb=hc(DB)
        else:
            termtbl.rows[ri].cells[ci*2].text=''; termtbl.rows[ri].cells[ci*2+1].text=''
doc.add_page_break()

# ── 封底 ─────────────────────────────────────────────────────────────────
tc_end = doc.add_table(rows=1,cols=1).cell(0,0)
bc2(tc_end, LB); cm(tc_end,350,350,350,350)
for line, sz, col, bold in [
    ('感谢您阅读本方案', 22, MB, True),
    (' ', 8, LB, False),
    ('我们期待与吉宝集团（Keppel Corporation）合作', 12, DB, False),
    ('共同建设东南亚下一代智慧港口', 12, DB, False),
    (' ', 10, LB, False),
    ('如有疑问，请联系：', 10, DB, False),
    ('淘景数科 — 业务发展部（总包方）', 10, MB, True),
    (' ', 8, LB, False),
    ('Keppel Corporation — Creating solutions for a sustainable future', 9, "666666", False),
]:
    p=tc_end.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name='Arial' if 'Keppel' in line or 'Creating' in line else 'SimSun'
        r.font.size=Pt(sz); r.font.color.rgb=hc(col)
        if bold: r.bold=True
    p.paragraph_format.space_after=Pt(3)

add_footer()

# ── 保存 ─────────────────────────────────────────────────────────────────
out = '/Users/tom/Desktop/智慧港口AI巡检解决方案_Keppel_V5_去AI味版.docx'
doc.save(out)
print('OK: ' + out)

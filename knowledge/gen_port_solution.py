# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Colors
DB  = "1F4E79"
MB  = "2E75B6"
LB  = "D6E4F0"
GR  = "595959"
ORG = "C55A11"
LOR = "FCE4D6"
WH  = "FFFFFF"
LG  = "F2F2F2"

def hc(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def sc(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    # remove existing shd
    for ch in tcPr:
        if ch.tag.endswith('}shd'):
            tcPr.remove(ch)
    tcPr.append(shd)

def bc(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top','left','bottom','right']:
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '1')
        el.set(qn('w:color'), color)
        # remove existing
        for ch in tcPr:
            if ch.tag.endswith('}w:' + side):
                tcPr.remove(ch)
        tcPr.append(el)

def cm(cell, t=120, b=120, l=200, r=200):
    tc = cell._tc
    # manually get or create w:tcMar
    tcMar = tc.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tc.append(tcMar)
    for side, v in [('top',t),('bottom',b),('left',l),('right',r)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(v))
        el.set(qn('w:type'), 'dxa')
        # remove existing side element
        for ch in list(tcMar):
            if ch.tag.endswith('}w:' + side):
                tcMar.remove(ch)
        tcMar.append(el)

def h1(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = hc(DB)
    r.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def h2(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = hc(MB)
    r.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def h3(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = hc(DB)
    r.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p

def para(t, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    if color:
        r.font.color.rgb = hc(color)
    else:
        r.font.color.rgb = hc(GR)
    p.paragraph_format.space_after = Pt(3)
    return p

def spacer():
    doc.add_paragraph()

def bullet(t, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(t)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = hc(GR)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def make_tbl(rows_data, col_pct, header=False, stripe=False):
    ncols = len(col_pct)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    pw = section.page_width - section.left_margin - section.right_margin
    for i, row_data in enumerate(rows_data):
        for j, text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.width = int(pw * col_pct[j] / 100)
            bc(cell)
            cm(cell)
            if header and i == 0:
                sc(cell, DB)
                for p2 in cell.paragraphs:
                    for r2 in p2.runs:
                        r2.font.color.rgb = hc(WH)
                        r2.bold = True
                        r2.font.size = Pt(8.5)
            elif stripe and i > 0 and i % 2 == 0:
                sc(cell, LG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # write text
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(text)
            r.font.name = 'Arial'
            r.font.size = Pt(8.5 if (header and i==0) else 8)
            if header and i == 0:
                r.bold = True
                r.font.color.rgb = hc(WH)
            else:
                r.font.color.rgb = hc(GR)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
    return tbl

def tip_box(lines):
    tbl = doc.add_table(rows=len(lines), cols=1)
    tbl.style = 'Table Grid'
    pw = section.page_width - section.left_margin - section.right_margin
    tbl.columns[0].width = pw
    for i, line in enumerate(lines):
        cell = tbl.rows[i].cells[0]
        cell.width = pw
        sc(cell, LOR)
        bc(cell, ORG)
        cm(cell, 80, 80, 150, 150)
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(line)
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        r.font.color.rgb = hc(DB)
    return tbl

# ═══════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════
spacer(); spacer(); spacer()
tc = doc.add_table(rows=1, cols=1).cell(0,0)
pw = section.page_width - section.left_margin - section.right_margin
tc.width = pw
sc(tc, DB)
bc(tc, DB)
cm(tc, 400, 400, 400, 400)
for line, size, color, bold in [
    ('SMART PORT', 28, 'AACCEE', True),
    ('智慧港口 AI 巡检与预警解决方案', 18, 'AACCEE', True),
    ('Smart Port AI Inspection & Early Warning Solution', 11, 'AACCEE', False),
    ('淘景数科 AI 视觉 + 大疆无人机 + 杭州旗晟机器人 + 智感时代无人船 + 智能道闸', 9, 'AACCEE', False),
]:
    p = tc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.color.rgb = hc(color)
    if bold:
        r.bold = True
    p.paragraph_format.space_after = Pt(4)

spacer(); spacer()
tc2 = doc.add_table(rows=1, cols=1).cell(0,0)
tc2.width = pw
sc(tc2, LB)
bc(tc2, MB)
cm(tc2, 200, 200, 300, 300)
for line, size, color, bold in [
    ('Keppel Corporation | 吉宝企业 | 淘景数科', 12, DB, True),
    ('高端一站式智慧港口解决方案 | 整体运行或拆卸组装灵活配置', 10, MB, False),
    ('2026 年 5 月 | 版本 3.0', 9, GR, False),
]:
    p = tc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.color.rgb = hc(color)
    if bold:
        r.bold = True
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 1
# ═══════════════════════════════════════════════════════
h1('第一章 核心定位与合作伙伴生态')
h2('1.1 整体定位：高端一站式解决方案')
para('本方案定位为高端一站式智慧港口解决方案，由淘景数科提供 AI 视觉算法与方案整合协同平台，联合大疆创新（无人机）、杭州旗晟（地面机器人）、北京智感时代（无人船）等顶尖供应商，形成完整生态链。')
spacer()
bullet('整体运行（Turnkey）：所有子系统（AI 视觉、边缘计算、机器人、无人机、道闸）无缝集成，开箱即用', 0)
bullet('拆卸组装（Modular）：各子系统支持独立部署，港口可按需选购单一模块（如仅采购 AI 视觉 + 道闸）', 0)
bullet('方案整合与协同：由淘景数科提供 T-Join 方案整合协同平台，实现多厂家设备统一接入与管理', 0)
bullet('灵活扩展：支持后续接入无人船、新型传感器等扩展模块，保护业主投资', 0)
spacer()
h2('1.2 核心合作伙伴与分工')
para('本方案联合国内顶尖供应商，形成完整生态链，各合作伙伴分工明确：')
spacer()
make_tbl([
    ['合作伙伴', '核心贡献', '覆盖能力'],
    ['淘景数科\n(Taojing Data Tech)', 'AI 视觉算法供应商\n方案整合与协同平台', '50+ 港口场景 AI 算法\n多厂家设备统一接入平台\n方案整合与协同指挥中心'],
    ['大疆创新\n(DJI)', '工业级无人机平台\n多传感器载荷', 'Matrice 350 RTK 无人机\nZenmuse H20T 四光融合\n可搭载清洗/灭火/空气质量传感器'],
    ['杭州旗晟\n(Hangzhou Qisheng)', '地面巡检机器人\n防爆/水陆双机型', '防爆巡检机器人（危品区）\n水陆行走机器人（冷藏区）\n全自动充电/换电基站'],
    ['北京智感时代\n(Zhigan Times)', '无人船补充方案\n割草打捞检测一体', '割草/打捞/水质检测三合一\n适用于内河港/运河港\n作为水上补充选项，按需配置'],
    ['Keppel Corporation\n(吉宝企业)', '整体方案集成商\n东南亚市场渠道', '系统集成与项目管理\n东南亚港口渠道与运营经验\n融资与全生命周期服务'],
], [20, 30, 50], header=True, stripe=True)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 2
# ═══════════════════════════════════════════════════════
h1('第二章 全港口场景 AI 视觉能力（淘景数科提供）')
h2('2.1 淘景数科 AI 视觉算法能力清单')
para('淘景数科作为本方案 AI 视觉算法核心供应商，提供覆盖 8 大业务领域、50+ 细分场景的深度学习算法集，支持边缘计算节点本地化部署。')
spacer()
make_tbl([
    ['算法分类', '算法名称（淘景数科）', '精度', '边缘部署'],
    ['集装箱识别类', 'YOLOv8/v9 实时目标检测\nOCR 箱号识别\n箱体缺陷分类 CNN', '≥ 99.5%', 'Jetson Nano\n≤ 30ms/帧'],
    ['安全监控类', '人员着装检测\n违章行为识别\n入侵检测\n火点检测', '≥ 99%', 'Jetson NX\n≤ 50ms/帧'],
    ['设备健康类', 'IR 温差成像设备过热检测\n链条磨损 CNN\n压力异常识别', '≥ 98%', 'Jetson AGX Orin\n≤ 100ms/帧'],
    ['环境监测类', '河面油污检测\n危品渗漏识别\n气体密度成像', '≥ 97%', 'Jetson NX\n≤ 200ms/帧'],
    ['运策优化类', '船位预测深度强化学习\nAGV 路径规划\n堆位优化 AI', '≥ 95%', 'Jetson AGX Orin\n≤ 500ms 推演'],
    ['多模态融合类', '可见光 + IR + 雾霭三模态融合\n5G + 视频多流协同', '≥ 99%', 'Jetson AGX Orin\n≤ 150ms/帧'],
], [20, 30, 20, 30], header=True, stripe=True)
spacer()
h2('2.2 方案整合与协同平台（淘景数科 T-Join 平台）')
para('淘景数科提供 T-Join 方案整合与协同平台，实现多厂家设备的统一接入、统一管理与统一告警：')
bullet('统一设备接入：支持大疆无人机、杭州旗晟机器人、智感时代无人船、海康/大华道闸等设备接入', 0)
bullet('统一 AI 算法管理：淘景数科 50+ 算法统一调度，支持按需加载/卸载', 0)
bullet('统一告警引擎：多级告警（紧急/重要/一般），支持微信/邮件/声光同步', 0)
bullet('统一 GIS 可视化：所有设备位置、告警位置、任务状态一张图展示', 0)
bullet('拆卸组装支持：各子系统 API 独立开放，支持第三方系统对接', 0)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 3
# ═══════════════════════════════════════════════════════
h1('第三章 地面巡检机器人能力（杭州旗晟提供）')
h2('3.1 杭州旗晟产品系列')
para('杭州旗晟是国内领先的特种巡检机器人制造商，为本方案提供三类地面巡检机器人：')
spacer()
make_tbl([
    ['产品类型', '适用场景', 'AI 能力', '防护等级', '续航'],
    ['防爆巡检机器人\n(QS-EX 系列)', '危品区、油品区\n化学品仓库', 'IR 温差成像 + 多气体传感器\n+ 可见光三融合', 'IP65\nEx d IIC T4\n防爆认证', '8h\n(自动换电)'],
    ['水陆行走机器人\n(QS-AMP 系列)', '冷藏集装箱区\n电缆通道', 'IR 摄像头缆损检测\n+ 插头接触状态识别\n+ 过滤器状态', 'IP54\n防溅水\n橡胶履带', '10h\n(自动充电)'],
    ['通用巡检机器人\n(QS-GEN 系列)', '堆场周界\n设施机房\n办公区域', '可见光 + IR 双光融合\n+ 超声波异常检测', 'IP54', '12h\n(自动充电)'],
], [20, 25, 25, 15, 15], header=True, stripe=True)
spacer()
h2('3.2 旗晟机器人巡检能力明细')
spacer()
make_tbl([
    ['业务领域', '机器人巡检项', '执行频率', '替代价值', '旗晟机型'],
    ['危品区', '渗漏/泄漏定期扫描\n气体浓度检测', '每 2h 一次', '代替人工进入高危区域\n降风险 100%', 'QS-EX 系列'],
    ['冷藏箱区', '电缆插头板板脱检测\n电缆损伤/老化监控\n过滤器状态检查', '每 4h 一次', '提前 30min 发现断电风险\n避免货损', 'QS-AMP 系列'],
    ['堆场', '火灾早期扫描\n堆场入侵巡查', '每 1h 一次\n实时', '提前 5-10min 发现初期火情\n实现 7×24h 无间断', 'QS-GEN 系列'],
    ['设施机房', '变压器/配电柜过热检测\n消防设施状态监控', '每 4h 一次', '提前 1-2h 发现过热苗头\n避免停电', 'QS-GEN 系列'],
    ['周界', '入侵巡查\n道闸异常监控', '实时', '实现大覆盖面周界无死角覆盖', 'QS-GEN 系列'],
], [15, 25, 15, 25, 20], header=True, stripe=True)
spacer()
h2('3.3 拆卸组装特性（旗晟机器人）')
para('杭州旗晟机器人支持灵活部署模式：')
bullet('整体运行模式：机器人 + 充电桩 + T-Join 平台全套交付，开箱即用', 0)
bullet('独立部署模式：机器人可独立运行，通过 API 对接业主现有系统（如 TOS、BMS）', 0)
bullet('混合部署模式：部分区域用旗晟机器人，部分区域保留人工，渐进式智能化', 0)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 4
# ═══════════════════════════════════════════════════════
h1('第四章 空中无人机巡检能力（大疆创新提供）')
h2('4.1 大疆无人机标准配置')
para('本方案采用大疆工业级无人机作为空中巡检平台，并支持搭载多种任务传感器：')
spacer()
make_tbl([
    ['配置项', '技术规格', '数量建议', '覆盖场景'],
    ['无人机机架\n(DJI Matrice 350 RTK)', 'IP55，抗风 7 级\n最大载重 2.7kg\n续航 55min', '每 5 万 TEU\n配置 1 架', '大覆盖面周界巡查\n岸吊顶部检查\n应急响应'],
    ['标准载荷\n(Zenmuse H20T)', '可见光 + IR\n+ 激光测距 + 雾霭\n四光融合', '通用 1 套/无人机', '全天候巡检\n夜间 IR 补偿\n3D 点云扫描'],
    ['航线基站\n(DJI Dock 2)', 'IP55，自动充电\n工作温度 -25~45℃\n支持远程调度', '每 5 万 TEU\n配置 1 个航点', '无人值守\n定期自动起降\n减少飞手 2 名'],
], [20, 35, 20, 25], header=True, stripe=True)
spacer()
h2('4.2 可扩展任务传感器（大疆生态）')
para('除标准视觉载荷外，大疆无人机支持搭载多种任务传感器，满足港口特定需求：')
spacer()
make_tbl([
    ['任务传感器', '功能描述', '适用场景', '大疆配套'],
    ['清洗喷头载荷', '高压水枪 + 清洗剂喷洒\n远程控制清洗', '岸吊玻璃污渍清洗\n摄像头镜头清洁\n太阳能板清洗', 'DJI Payload SDK\n第三方定制'],
    ['灭火弹/水弹载荷', '灭火弹远程投放\n或水弹精准喷射', '危品区初期火情处置\n堆场火点快速压制\n高空灭火支援', 'DJI Payload SDK\n应急专用载荷'],
    ['空气质量传感器', 'PM2.5/PM10\nSO2/NO2/CO 多气体\n实时回传浓度场', '港区空气质量监测\n危品区气体泄漏预警\n环保合规监测', 'DJI Payload SDK\n大气监测载荷'],
    ['多光谱水质仪', '叶绿素/浊度/COD\n多光谱成像', '河面水质监测\n油污扩散追踪\n水产养殖区监测', 'DJI Payload SDK\n水质监测载荷'],
    ['喊话/照明一体', '远程喊话 + 强光照明\n夜间应急指挥', '周界入侵喊话警告\n夜间应急照明\n搜救现场指挥', 'DJI Speaker\n+ Beacon'],
], [20, 35, 25, 20], header=True, stripe=True)
spacer()
h2('4.3 无人机巡检能力明细')
spacer()
make_tbl([
    ['业务领域', '无人机巡检项', '执行频率', '大疆机型', '替代价值'],
    ['岸吊顶部', '岸吊链条/滑轮等高吊设备检查\n清洗喷头定期清洁摄像头', '每周 1 次', 'Matrice 350 RTK\n+ H20T', '代替人工高空作业\n降风险 100%'],
    ['堆场大覆盖面', '堆场全景覆盖盘点\nIR + 可见光双光 AI 识别', '每天 2 次', 'Matrice 350 RTK\n+ Dock 2', '效率提升 15x\n覆盖率 100%'],
    ['周界大覆盖面', '周界入侵巡查\n喊话警告 + 夜间 IR', '每天 4 次', 'Matrice 350 RTK\n+ Speaker', '实现大覆盖面周界无死角'],
    ['航道与水域', '河面监控（河面油污）\n水质多光谱监测', '每天 1 次', 'Matrice 350 RTK\n+ 水质载荷', '提前 15min 发现河面河污'],
    ['危品区', '危品区全景风险评估\n气体传感器数据采集\n灭火弹应急储备', '每周 2 次\n应急即时', 'Matrice 350 RTK\n+ 气体载荷\n+ 灭火载荷', '提供危品区立体风险扫描\n初期火情快速压制'],
    ['应急响应', '事故现场快速报到\n喊话/照明/灭火多任务', '按需', 'Matrice 350 RTK\n全载荷', '应急响应时间从 30min → 5min'],
], [15, 25, 15, 20, 25], header=True, stripe=True)
spacer()
h2('4.4 拆卸组装特性（大疆无人机）')
para('大疆无人机系统支持灵活配置：')
bullet('整体运行：无人机 + Dock 2 基站 + T-Join 平台全套交付，自动起降无需飞手', 0)
bullet('独立部署：仅采购无人机+遥控器，由业主现有团队操作，对接 T-Join API', 0)
bullet('载荷自选：业主可按需选购标准载荷或扩展传感器（清洗/灭火/空气质量等）', 0)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 5
# ═══════════════════════════════════════════════════════
h1('第五章 无人船补充方案（北京智感时代提供）')
h2('5.1 智感时代产品：割草打捞检测一体无人船')
para('针对内河港、运河港等水域场景，本方案将北京智感时代研发的割草打捞检测一体无人船作为水上补充选项，按需配置。')
spacer()
make_tbl([
    ['配置项', '技术规格（智感时代）', '适用场景', '补充价值'],
    ['无人船平台\n(ZG-WS 系列)', '体长 3.2m，续航 6h\n最大航速 3m/s\nIP65 防水', '内河港水面巡查\n运河港航道维护\n港区水域保洁', '代替人工船艇\n降人工成本 80%'],
    ['割草打捞一体\n载荷', '自动割草 + 打捞\n垃圾仓容量 200L\n割草深度 0.5-2m', '港区水域水草清理\n水面漂浮物打捞\n航道疏浚辅助', '水域保洁效率提升 10x\n减少人工船艇 3 艘'],
    ['水质检测载荷', '多光谱水质仪\nCOD/浊度/叶绿素\n实时回传浓度场', '河面水质监测\n油污扩散追踪\n环保合规报告', '水质监测效率提升 20x\n数据自动上传'],
    ['AI 视觉补充\n巡检', '可见光 + IR 双光\n搭载淘景数科 AI 算法', '水面异常检测\n危品区水域监控\n河面油污识别', '与无人机/机器人\n形成水陆空一体'],
], [20, 35, 25, 20], header=True, stripe=True)
spacer()
h2('5.2 无人船应用场景明细')
spacer()
make_tbl([
    ['业务领域', '无人船作业项', '执行频率', '智感时代机型'],
    ['港区水域', '水面割草 + 打捞\n垃圾清运', '每 3 天一次', 'ZG-WS 系列\n割草打捞一体'],
    ['航道维护', '航道疏浚辅助\n水面漂浮物清理', '每 7 天一次', 'ZG-WS 系列\n大仓容量型'],
    ['水质监测', '多点位水质采样\n浓度场实时回传', '每天 1 次', 'ZG-WS 系列\n水质检测型'],
    ['应急补充', '河面油污回收辅助\n水面异常喊话警告', '按需', 'ZG-WS 系列\n+ 喊话载荷'],
], [25, 35, 20, 20], header=True, stripe=True)
spacer()
h2('5.3 拆卸组装特性（智感时代无人船）')
para('北京智感时代无人船支持灵活部署：')
bullet('补充选项：不作为标配，仅在内河港/运河港场景中作为水上补充选项推荐', 0)
bullet('独立运行：无人船可独立运行，通过 API 对接 T-Join 平台或业主现有系统', 0)
bullet('载荷自选：业主可按需选购割草型/打捞型/检测型，或三合一全能型', 0)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 6
# ═══════════════════════════════════════════════════════
h1('第六章 港口道闸设备（车辆出入自动化）')
h2('6.1 道闸系统功能清单')
para('本方案集成智能道闸系统，实现港口车辆出入的自动化登记、审批、记录与联动：')
spacer()
make_tbl([
    ['功能模块', '技术实现', '覆盖场景', '价值'],
    ['车牌自动识别', '多国车牌 OCR\n支持东南亚各国车牌\n夜间 IR 补偿', '闸口入港/出港\n停车场出入口', '通过时间 ≤ 25秒\n无人化率 ≥ 90%'],
    ['箱号 OCR 识别', 'OCR 箱号识别\n精度 ≥ 99.7%\n与 TOS 系统对接', '闸口通道\n集装箱卡车查验', '自动完成箱号校验\n减少人工录入错误'],
    ['司机身份认证', '人脸识别 + 身份证 OCR\n与派单系统联动', '闸口通行\n危品区入口', '无感通行\n未授权人员拦截率 100%'],
    ['自动审批流程', 'T-Join 平台审批引擎\n支持多级审批\n微信/邮件通知', '外协车辆准入\n危品车辆审批\n夜间放行审批', '审批时间从 30min → 3min\n审批记录可追溯'],
    ['车辆外观损伤检测', 'AI 视觉车辆外观损伤 CNN\n进港时自动记录', '闸口通道\n车辆交接验收', '减少车辆损伤纠纷\n责任界定清晰'],
    ['道闸异常监控', '道闸运行状态 IR 监测\n防砸车传感器\n故障自诊断', '所有道闸点位', '道闸故障率降低 80%\n砸车事故零发生'],
], [20, 35, 25, 20], header=True, stripe=True)
spacer()
h2('6.2 道闸与 AI 巡检系统协同')
para('道闸系统作为港口车辆管理的第一道防线，与 AI 巡检系统形成协同：')
bullet('道闸识别异常 → 自动触发附近 CCTV 摄像头多角度确认 → 告警推送至 T-Join 平台', 0)
bullet('危品车辆过闸 → 自动触发危品区机器人/无人机加强巡查频次', 0)
bullet('未授权车辆强行过闸 → 自动触发周界入侵告警 + 无人机喊话警告', 0)
bullet('道闸审批数据与 TOS 对接，实现车辆到港预判，优化堆场作业计划', 0)
spacer()
h2('6.3 拆卸组装特性（道闸系统）')
para('道闸系统支持灵活部署模式：')
bullet('整体运行：道闸 + T-Join 平台 + AI 视觉识别全套交付，自动审批无需人工干预', 0)
bullet('独立部署：仅采购道闸硬件 + 基础车牌识别，对接业主现有审批系统', 0)
bullet('混合部署：部分闸口用智能道闸，部分闸口保留人工，渐进式智能化', 0)
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 7
# ═══════════════════════════════════════════════════════
h1('第七章 多能力协同体系与不可或缺性')
h2('7.1 淘景数科 T-Join 平台协同架构')
para('淘景数科 T-Join 平台作为本方案的核心协同中枢，实现多厂家、多类型设备的统一接入与协同：')
spacer()
make_tbl([
    ['协同层级', '接入设备/系统', '淘景数科 T-Join 功能', '协同价值'],
    ['AI 视觉层', '固定 CCTV（海康/大华）\n淘景数科 50+ 算法', '算法统一调度\n多路视频流并行处理\n边缘-云端协同推理', 'AI 能力整体输出\n单点预警全域联动'],
    ['机器人层', '杭州旗晟机器人\n（防爆/水陆/通用）', '机器人任务派单引擎\n多机协同路径规划\n充电/换电自动调度', '多机器人协同作业\n无死角覆盖'],
    ['无人机层', '大疆无人机\n扩展传感器载荷', '航线自动规划\n多传感器数据采集\n自动起降任务调度', '空中补充覆盖\n应急响应 ≤ 5min'],
    ['无人船层\n(补充)', '北京智感时代无人船\n(按需配置）', '水域任务规划\n与无人机/机器人数据融合\n环保合规报告自动生成', '水陆空一体\n内河港完整方案'],
    ['道闸层', '智能道闸系统\n车牌/箱号/OCR', '车辆准入审批引擎\n与 TOS 系统对接\n道闸异常联动告警', '车辆管理第一道防线\n与巡检系统联动'],
], [18, 25, 30, 27], header=True, stripe=True)
spacer()
h2('7.2 端到端巡检闭环流程（T-Join 平台驱动）')
spacer()
make_tbl([
    ['流程环节', '执行单元', '淘景数科 T-Join 支撑', '时效指标'],
    ['1. 异常发现', '固定 CCTV / 旗晟机器人\n/ 大疆无人机', 'AI 算法自动检测\n边缘推理 + 云端复核', '实时（≤ 100ms）'],
    ['2. 多级告警', 'T-Join 平台', '告警级别分层：紧急/重要/一般\n微信/邮件/声光同步', '告警延迟 ≤ 2s'],
    ['3. 任务派单', 'T-Join 任务引擎', '自动派单至最近机器人/无人机\n支持人工干预派单', '派单延迟 ≤ 10s'],
    ['4. 现场处置', '旗晟机器人 / 大疆无人机\n/ 智感时代无人船（补充）', '处置结果拍照上传\n如需支援自动触发协同单元', '处置时效依级别'],
    ['5. 复核归档', 'T-Join AI 自动复核', '处置前后对比图 AI 自动比对\n形成事件档案，支持回溯', '复核延迟 ≤ 60s'],
], [18, 25, 30, 27], header=True, stripe=True)
spacer()
h2('7.3 各能力单元不可或缺性论证')
para('本方案强调"单维能力有限，多维协同无限"，各能力单元缺一不可：')
bullet('无淘景数科 AI 算法：则所有视频需人工查看，漏检率 ≥ 95%，方案无核心价值', 0)
bullet('无杭州旗晟机器人：则危品区、冷藏箱区等高危/狭隘区域无法巡检，人工风险极高', 0)
bullet('无大疆无人机：则岸吊顶部、大覆盖面周界等高空/大区域无法快速巡查，漏检率 ≥ 70%', 0)
bullet('无北京智感时代无人船：则内河港/运河港水域场景无法覆盖，方案不完整（仅作为补充选项）', 0)
bullet('无智能道闸：则车辆出入无法自动化，审批效率低下，与 TOS 对接断层', 0)
bullet('无 T-Join 协同平台：则各子系统孤立运行，无法形成协同效应，整体效能大打折扣', 0)
spacer()
tip_box([
    '核心结论：本方案 淘景 AI + 旗晟机器人 + 大疆无人机 + 智感无人船（补充）+ 智能道闸 多位一体，',
    '单维能力无法独立达成 全覆盖、无死角、提前预警、自动处置 目标。协同价值：1+1+1+1+1 > 5。',
])
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# CHAPTER 8
# ═══════════════════════════════════════════════════════
h1('第八章 高端一站式方案价值与实施路径')
h2('8.1 高端一站式解决方案核心优势')
para('本方案作为高端一站式智慧港口解决方案，具备以下核心优势：')
spacer()
make_tbl([
    ['优势维度', '详细说明', '客户价值'],
    ['整体运行\n（Turnkey）', '所有子系统由淘景数科 T-Join 平台统一集成\n开箱即用，业主无需对接多个供应商', '项目交付周期缩短 50%\n责任边界清晰\n终身维护统一接口'],
    ['拆卸组装\n（Modular）', '各子系统支持独立部署\nAPI 独立开放，支持第三方系统对接\n业主可按需选购单一模块', '保护既有投资\n渐进式智能化\n降低初期 Capex 压力'],
    ['方案整合\n与协同', '淘景数科提供方案整合与协同平台\n多厂家设备统一接入、统一管理、统一告警', '避免多厂家推诿\n系统互联互通\n数据融合分析'],
    ['灵活扩展', '支持后续接入新型传感器、新型机器人/无人机\n智感时代无人船作为补充选项按需配置', '技术迭代无忧\n保护业主长远投资\n方案可持续演进'],
], [22, 38, 40], header=True, stripe=True)
spacer()
h2('8.2 设备清单与投资测算（基准：100 万 TEU/年）')
spacer()
make_tbl([
    ['系统/设备', '供应商', '数量（100万TEU）', '投资估算（亿元）'],
    ['淘景数科 AI 视觉系统\n(T-Join 平台 + 50+ 算法）', '淘景数科', '1 套', '约 0.8-1.2 亿'],
    ['杭州旗晟巡检机器人\n(防爆+水陆+通用）', '杭州旗晟', '8-12 台\n+ 3-5 座充电桩', '约 0.6-1.0 亿'],
    ['大疆无人机系统\n(含 Dock 2 基站 + 扩展载荷）', '大疆创新', '1-2 架\n+ 1 个航点', '约 0.3-0.5 亿'],
    ['北京智感时代无人船\n(补充选项，按需）', '北京智感时代', '1-2 艘\n(仅内河港配置）', '约 0.1-0.2 亿\n(可选）'],
    ['智能道闸系统\n(车牌+OCR+审批）', '海康/大华/科拓', '10-15 通道', '约 0.2-0.4 亿'],
    ['边缘计算硬件\n(NVIDIA Jetson 系列）', '英伟达 + 淘景集成', '12-20 台', '约 0.3-0.6 亿'],
    ['合计总投资\n(不含道闸基础建设）', '—', '—', '约 2.3-3.9 亿\n(不含智感无人船）'],
], [28, 22, 22, 28], header=True, stripe=True)
spacer()
h2('8.3 年度收益与投资回收期')
spacer()
make_tbl([
    ['收益/成本项目', '年度金额（人民币）', '说明'],
    ['作业效率提升带来的产能收益', '约 2.5-4 亿元/年', '吞吐量提升 20-30%，新增收入按码头收入测算'],
    ['人力成本节省', '约 1.2-1.8 亿元/年', '减少约 1000 人，按年人力成本 12-18 万元测算'],
    ['设备维护成本降低', '约 3000-5000 万元/年', '预测性维护减少非计划停机，维护费用降低 25-35%'],
    ['AI 巡检替代人工降本', '约 8000万-1.2 亿元/年', '代替高危/重复人工巡检，降本 70%+'],
    ['合计年度收益', '约 5.3-7.7 亿元/年', '—'],
    ['静态投资回收期', '约 0.4-0.7 年\n(约 5-9 个月）', '高端一站式方案 ROI 极高\n含设备折旧后实际回收期约 1-1.5 年'],
], [35, 30, 35], header=True, stripe=True)
spacer()
h2('8.4 实施路径（4 年分阶段）')
spacer()
make_tbl([
    ['阶段', '时间', '建设内容', '投资估算'],
    ['第一阶段\n基础建设', 'Year 1-2', '淘景数科 T-Join 平台部署\n道闸系统安装与对接\n边缘计算节点部署（AI 视觉试点）', '约 0.8-1.5 亿'],
    ['第二阶段\n机器人与道闸全面上线', 'Year 2-3', '杭州旗晟机器人部署\n道闸系统全覆盖\n淘景 AI 算法全量加载', '约 0.6-1.0 亿'],
    ['第三阶段\n无人机与协同优化', 'Year 3-4', '大疆无人机系统上线\nT-Join 协同平台全面优化\n智感时代无人船补充（按需）', '约 0.5-0.8 亿'],
    ['第四阶段\n持续优化', 'Year 4+', 'AI 模型持续优化\n新型传感器/机器人接入\n系统运维与升级', '约 0.2-0.3 亿/年'],
], [15, 18, 45, 22], header=True, stripe=True)
spacer()
h2('8.5 核心结论与建议')
para('通过本方案的系统设计，得出以下核心结论：')
bullet('高端一站式方案是出海核心竞争力：整体运行（Turnkey）+ 拆卸组装（Modular）双重灵活性，满足不同港口的差异化需求', 0)
bullet('淘景数科作为 AI 视觉与方案整合核心，是整体方案的技术基石，不可或缺', 0)
bullet('大疆无人机 + 杭州旗晟机器人 + 智能道闸形成"空地一体"巡检体系，覆盖港口全场景', 0)
bullet('北京智感时代无人船作为水上补充选项，完善内河港/运河港场景覆盖，提升方案完整性', 0)
bullet('投资回收期极短（5-9 个月），战略价值极高，是东南亚港口智慧化升级的首选方案', 0)
spacer()
tip_box([
    '近期（6 个月内）：与淘景数科、杭州旗晟、大疆签订战略合作协议，锁定供应链',
    '中期（1-2 年）：锁定 1-2 个东南亚港口试点，优先推进道闸 + AI 视觉 + 机器人基础建设',
    '长期（3-5 年）：建立智慧港口运营联盟（Keppel + 淘景 + 旗晟 + 大疆），输出标准化方案至全球 50+ 港口',
])
spacer()

# ═══════════════════════════════════════════════════════
# END BLOCK
# ═══════════════════════════════════════════════════════
tc3 = doc.add_table(rows=1, cols=1).cell(0,0)
tc3.width = pw
sc(tc3, LB)
bc(tc3, MB)
cm(tc3, 200, 200, 300, 300)
for line, size, color, bold, ital in [
    ('如需进一步了解本方案的技术细节、投资测算或合作模式，请联系 Keppel Corporation 商务拓展团队。', 10, DB, False, False),
    ('Keppel Corporation — Creating solutions for a sustainable future', 9, MB, False, True),
]:
    p = tc3.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.color.rgb = hc(color)
    if bold:
        r.bold = True
    if ital:
        r.italic = True
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════
out = '/Users/tom/.qclaw/workspace/智慧港口AI巡检解决方案_Keppel_淘景大疆旗晟_2026V3.docx'
doc.save(out)
print('OK: ' + out)
